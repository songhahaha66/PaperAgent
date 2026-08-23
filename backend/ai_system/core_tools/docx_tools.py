import logging
import asyncio
import os
import json
import shutil
import subprocess
import zipfile
from pathlib import Path
from typing import Optional, List

from .docx_images import extract_docx_images, format_image_inventory, inventory_docx_images
from .docx_styles import (
    apply_body_paragraph_format,
    apply_body_run_format,
    compare_style_fingerprints,
    extract_style_fingerprint,
    format_style_comparison,
)

logger = logging.getLogger(__name__)

SKILL_DIR = Path(__file__).resolve().parent.parent / "docx_skill"
SCRIPTS_DIR = SKILL_DIR / "scripts"


def _node_path() -> str:
    return os.environ.get("NODE_PATH", "") or subprocess.check_output(
        ["npm", "root", "-g"], text=True
    ).strip()


class DocxTools:
    def __init__(self, workspace_dir: str, stream_manager=None):
        self.workspace_dir = Path(workspace_dir).resolve()
        self.stream_manager = stream_manager
        self.document_path = self.workspace_dir / "paper.docx"

    def _notify_file_changed(self):
        if not self.stream_manager:
            return
        try:
            asyncio.create_task(
                self.stream_manager.send_json_block("file_changed", "paper.docx")
            )
        except Exception as e:
            logger.warning(f"Failed to send file_changed notification: {e}")

    def _has_template(self) -> bool:
        return (self.workspace_dir / ".system" / "_template_original.docx").exists()

    async def create_docx(self, js_code: str, filename: str = "paper.docx") -> str:
        """
        用 docx-js 的 JavaScript 代码创建 .docx 文件。

        AI 应生成完整的 JS 脚本，使用 require('docx') 和 require('fs')
        来构建文档并写入文件。脚本会在工作空间目录下执行。

        ⚠️ 当存在模板时，不允许覆盖 paper.docx，请使用 write_to_template 工具。

        Args:
            js_code: 完整的 Node.js 脚本，使用 docx-js 创建文档。
                     脚本中应使用 process.env.OUTPUT_PATH 获取输出路径。
            filename: 输出文件名（默认 paper.docx）

        Returns:
            执行结果，成功时包含文件路径，失败时包含错误信息
        """
        if filename == "paper.docx" and self._has_template():
            return (
                "Error: 当前工作空间存在模板文件，禁止用 create_docx 覆盖 paper.docx。\n"
                "请使用 write_to_template 工具在模板基础上填充内容。\n"
                "如果需要创建其他文件，请指定不同的 filename 参数。"
            )
        output_path = self.workspace_dir / filename
        js_file = self.workspace_dir / ".system" / "_docx_gen.js"
        js_file.parent.mkdir(parents=True, exist_ok=True)

        wrapper = (
            f"process.env.OUTPUT_PATH = {str(output_path)!r};\n"
            f"{js_code}\n"
        )
        js_file.write_text(wrapper, encoding="utf-8")

        try:
            env = os.environ.copy()
            env["NODE_PATH"] = _node_path()
            proc = await asyncio.create_subprocess_exec(
                "node", str(js_file),
                cwd=str(self.workspace_dir),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
                env=env,
            )
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=30)

            if proc.returncode != 0:
                err = stderr.decode(errors="replace").strip()
                logger.error(f"docx-js 执行失败: {err}")
                return f"Error: JS 执行失败 (exit {proc.returncode}):\n{err}"

            if not output_path.exists():
                return "Error: JS 脚本执行成功但未生成文件，请确保脚本写入了 process.env.OUTPUT_PATH"

            validate_result = await self._validate(output_path)
            image_result = self._ensure_workspace_images_in_docx(output_path)
            self._notify_file_changed()

            size_kb = output_path.stat().st_size / 1024
            result = f"✅ {filename} 创建成功 ({size_kb:.1f} KB)"
            if stdout.decode().strip():
                result += f"\n{stdout.decode().strip()}"
            if validate_result:
                result += f"\n{validate_result}"
            if image_result:
                result += f"\n{image_result}"
            return result

        except asyncio.TimeoutError:
            return "Error: JS 脚本执行超时（30秒限制）"
        except Exception as e:
            logger.error(f"create_docx 失败: {e}", exc_info=True)
            return f"Error: {e}"

    def _ensure_workspace_images_in_docx(self, docx_path: Path) -> str:
        """
        Add user-visible workspace images to the generated Word file when the
        LLM-created docx forgot to include them.

        This is a conservative fallback: if the docx already contains embedded
        media, it does nothing. Images are discovered from outputs/ first, then
        manifest-registered artifacts, then latest run artifacts.
        """
        try:
            if self._docx_has_images(docx_path):
                return ""

            images = self._find_workspace_images()
            if not images:
                return ""

            from docx import Document as PythonDocxDocument
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from PIL import Image

            doc = PythonDocxDocument(str(docx_path))
            doc.add_page_break()
            heading = doc.add_paragraph()
            heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
            heading_run = heading.add_run("附图")
            heading_run.bold = True
            heading_run.font.size = Pt(16)

            for index, image_path in enumerate(images, start=1):
                caption = doc.add_paragraph()
                caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
                caption.add_run(f"图{index}：{image_path.stem.replace('_', ' ')}").bold = True

                max_width_inches = 6.0
                width_inches = max_width_inches
                try:
                    with Image.open(image_path) as img:
                        px_width, px_height = img.size
                    if px_width and px_height:
                        # Keep the rendered image within a normal document page.
                        width_inches = min(max_width_inches, max(3.0, px_width / 700))
                except Exception:
                    width_inches = max_width_inches

                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = para.add_run()
                run.add_picture(str(image_path), width=Inches(width_inches))

            doc.save(str(docx_path))
            return f"已自动插入 {len(images)} 张工作区图片到 {docx_path.name}"

        except ImportError as e:
            logger.warning("无法自动插入图片，缺少依赖: %s", e)
            return "⚠️ 未能自动插入图片：缺少 python-docx 或 Pillow"
        except Exception as e:
            logger.warning("自动插入工作区图片失败: %s", e, exc_info=True)
            return f"⚠️ 自动插入图片失败: {e}"

    def _docx_has_images(self, docx_path: Path) -> bool:
        try:
            with zipfile.ZipFile(docx_path) as zf:
                return any(name.startswith("word/media/") for name in zf.namelist())
        except Exception:
            return False

    def _find_workspace_images(self) -> List[Path]:
        image_exts = {".png", ".jpg", ".jpeg", ".gif", ".bmp", ".webp"}
        candidates: List[Path] = []

        def add(path: Path):
            try:
                resolved = path.resolve()
                if (
                    resolved.is_file()
                    and resolved.suffix.lower() in image_exts
                    and str(resolved).startswith(str(self.workspace_dir))
                    and resolved not in candidates
                ):
                    candidates.append(resolved)
            except Exception:
                return

        outputs_dir = self.workspace_dir / "outputs"
        if outputs_dir.exists():
            for path in sorted(outputs_dir.rglob("*")):
                add(path)

        manifest_path = self.workspace_dir / "manifest.json"
        if manifest_path.exists():
            try:
                manifest = json.loads(manifest_path.read_text(encoding="utf-8"))
                files = manifest.get("files", {})
                for rel_path, meta in files.items():
                    if isinstance(meta, dict) and meta.get("visibility") == "user":
                        add(self.workspace_dir / rel_path)
            except Exception as e:
                logger.debug("读取 manifest 图片失败: %s", e)

        if not candidates:
            runs_dir = self.workspace_dir / "runs"
            if runs_dir.exists():
                artifact_dirs = sorted(
                    (p for p in runs_dir.glob("run_*/artifacts") if p.is_dir()),
                    key=lambda p: p.stat().st_mtime,
                    reverse=True,
                )
                for artifact_dir in artifact_dirs[:1]:
                    for path in sorted(artifact_dir.rglob("*")):
                        add(path)

        return candidates

    async def read_docx(self, filename: str = "paper.docx") -> str:
        """
        读取 .docx 文件内容为纯文本。

        Args:
            filename: 要读取的文件名（默认 paper.docx）

        Returns:
            文档的纯文本内容，或错误信息
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"

        text = ""
        try:
            pandoc = shutil.which("pandoc")
            if pandoc:
                proc = await asyncio.create_subprocess_exec(
                    pandoc, str(file_path), "-t", "plain", "--wrap=none",
                    stdout=asyncio.subprocess.PIPE,
                    stderr=asyncio.subprocess.PIPE,
                )
                stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=15)
                if proc.returncode == 0:
                    text = stdout.decode(errors="replace")
        except Exception as e:
            logger.warning(f"pandoc 读取失败，回退到 python-docx: {e}")

        if not text:
            try:
                from docx import Document as PythonDocxDocument
                doc = PythonDocxDocument(str(file_path))
                paragraphs = [p.text for p in doc.paragraphs]
                text = "\n".join(paragraphs)
            except ImportError:
                text = ""

        if text:
            images = inventory_docx_images(file_path)
            if images:
                text = text.rstrip() + "\n\n" + format_image_inventory(images, title="嵌入图片")
            return text

        try:
            unpack_dir = self.workspace_dir / ".system" / "_docx_unpacked"
            if unpack_dir.exists():
                shutil.rmtree(unpack_dir)
            proc = await asyncio.create_subprocess_exec(
                "python3", str(SCRIPTS_DIR / "office" / "unpack.py"),
                str(file_path), str(unpack_dir),
                cwd=str(SCRIPTS_DIR / "office"),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            await asyncio.wait_for(proc.communicate(), timeout=15)
            doc_xml = unpack_dir / "word" / "document.xml"
            if doc_xml.exists():
                import re
                content = doc_xml.read_text(encoding="utf-8")
                text = re.sub(r"<[^>]+>", "", content)
                text = re.sub(r"\s+", " ", text).strip()
                images = inventory_docx_images(file_path)
                if images:
                    text = text.rstrip() + "\n\n" + format_image_inventory(images, title="嵌入图片")
                return text
        except Exception as e:
            logger.warning(f"unpack 读取失败: {e}")

        return "Error: 无法读取文件（pandoc、python-docx、unpack 均不可用）"

    async def edit_docx(self, operations: str, filename: str = "paper.docx") -> str:
        """
        通过解包 → 编辑 XML → 重打包来编辑已有 .docx 文件。

        Args:
            operations: 自然语言描述的编辑操作（由 AI 解析执行）。
                       当前实现先解包，返回 XML 结构供 AI 分析。
            filename: 要编辑的文件名

        Returns:
            解包后的文件结构和关键 XML 内容
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"

        unpack_dir = self.workspace_dir / ".system" / "_docx_edit"
        if unpack_dir.exists():
            shutil.rmtree(unpack_dir)

        try:
            proc = await asyncio.create_subprocess_exec(
                "python3", str(SCRIPTS_DIR / "office" / "unpack.py"),
                str(file_path), str(unpack_dir),
                cwd=str(SCRIPTS_DIR / "office"),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=15)

            if proc.returncode != 0:
                return f"Error: 解包失败: {stderr.decode(errors='replace')}"

            files = []
            for root, _, fnames in os.walk(unpack_dir):
                for fn in fnames:
                    rel = os.path.relpath(os.path.join(root, fn), unpack_dir)
                    files.append(rel)

            doc_xml = unpack_dir / "word" / "document.xml"
            preview = ""
            if doc_xml.exists():
                content = doc_xml.read_text(encoding="utf-8")
                preview = content[:3000]
                if len(content) > 3000:
                    preview += f"\n... (共 {len(content)} 字符)"

            return (
                f"✅ 已解包到 .system/_docx_edit/\n"
                f"文件列表:\n" + "\n".join(f"  {f}" for f in sorted(files)) +
                f"\n\n--- document.xml 预览 ---\n{preview}"
            )

        except Exception as e:
            return f"Error: 编辑操作失败: {e}"

    async def repack_docx(self, filename: str = "paper.docx") -> str:
        """
        将编辑后的 XML 重新打包为 .docx 文件。

        前提：已通过 edit_docx 解包，并在 .system/_docx_edit/ 中修改了 XML 文件。

        Args:
            filename: 输出文件名

        Returns:
            打包结果
        """
        unpack_dir = self.workspace_dir / ".system" / "_docx_edit"
        if not unpack_dir.exists():
            return "Error: 未找到解包目录，请先调用 edit_docx"

        output_path = self.workspace_dir / filename
        original = self.workspace_dir / filename

        try:
            args = [
                "python3", str(SCRIPTS_DIR / "office" / "pack.py"),
                str(unpack_dir), str(output_path),
            ]
            if original.exists():
                args.extend(["--original", str(original)])

            proc = await asyncio.create_subprocess_exec(
                *args,
                cwd=str(SCRIPTS_DIR / "office"),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=15)

            if proc.returncode != 0:
                return f"Error: 打包失败: {stderr.decode(errors='replace')}"

            self._notify_file_changed()
            size_kb = output_path.stat().st_size / 1024
            return f"✅ {filename} 重新打包成功 ({size_kb:.1f} KB)"

        except Exception as e:
            return f"Error: 重新打包失败: {e}"

    async def _validate(self, file_path: Path) -> str:
        validate_script = SCRIPTS_DIR / "office" / "validate.py"
        if not validate_script.exists():
            return ""
        try:
            proc = await asyncio.create_subprocess_exec(
                "python3", str(validate_script), str(file_path),
                cwd=str(SCRIPTS_DIR / "office"),
                stdout=asyncio.subprocess.PIPE,
                stderr=asyncio.subprocess.PIPE,
            )
            stdout, stderr = await asyncio.wait_for(proc.communicate(), timeout=10)
            output = stdout.decode(errors="replace").strip()
            if proc.returncode != 0:
                err = stderr.decode(errors="replace").strip()
                return f"⚠️ 验证警告: {err or output}"
            if output:
                return f"验证: {output}"
            return ""
        except Exception:
            return ""

    async def write_to_template(self, anchor_text: str, content: str,
                                position: str = "after",
                                style: str = "",
                                filename: str = "paper.docx") -> str:
        """
        在模板文档中定位 anchor_text 所在段落，然后插入内容。
        保留模板的所有原有格式和样式。

        Args:
            anchor_text: 用于定位的文本（匹配包含该文本的段落）
            content: 要插入的内容。多段落用 \\n\\n 分隔。
                     如果以 ``` 开头则视为代码块，自动应用等宽字体。
            position: "after"=在锚点段后插入, "before"=在锚点段前插入,
                      "replace"=替换锚点段落内容
            style: 可选，段落样式名（如 "Normal", "List Paragraph"）。
                   留空则继承模板默认样式。
            filename: 目标文件名（默认 paper.docx）

        Returns:
            操作结果
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"

        try:
            from docx import Document as PythonDocxDocument
            from docx.shared import Pt, Cm
            from docx.oxml.ns import qn
            from docx.oxml import OxmlElement
            import copy

            doc = PythonDocxDocument(str(file_path))

            def _style_name(para) -> str:
                return (para.style.name if para.style else "").strip().lower()

            def _is_toc_para(para) -> bool:
                style_name = _style_name(para)
                return style_name.startswith("toc") or "目录" in style_name

            def _matches(para, *, exact: bool) -> bool:
                text = para.text.strip()
                target = anchor_text.strip()
                if not text or _is_toc_para(para):
                    return False
                return text == target if exact else target in text

            anchor_idx = None
            for exact in (True, False):
                for i, para in enumerate(doc.paragraphs):
                    if _matches(para, exact=exact):
                        anchor_idx = i
                        break
                if anchor_idx is not None:
                    break

            if anchor_idx is None:
                available = [
                    f"[{i}] {p.text[:60]}"
                    for i, p in enumerate(doc.paragraphs)
                    if p.text.strip() and not _is_toc_para(p)
                ][:30]
                return (
                    f"Error: 未找到包含 \"{anchor_text}\" 的段落。\n"
                    f"可用段落（前30个非空）:\n" + "\n".join(available)
                )

            is_code = content.startswith("```")
            if is_code:
                lines = content.strip().split("\n")
                if lines[0].startswith("```"):
                    lines = lines[1:]
                if lines and lines[-1].strip() == "```":
                    lines = lines[:-1]
                paragraphs_text = ["\n".join(lines)]
            else:
                paragraphs_text = [p.strip() for p in content.split("\n\n") if p.strip()]

            if not paragraphs_text:
                return "Error: content 为空"

            anchor_para = doc.paragraphs[anchor_idx]
            anchor_style = _style_name(anchor_para)
            if position == "replace" and anchor_style.startswith("heading"):
                return (
                    "Error: 禁止替换模板标题段落。"
                    "请保留标题骨架，并使用 position='after' 在标题或占位说明后填充内容。"
                )

            body_style = style or "Normal"

            def _styled_paragraph(text: str, after_element=None, insert_first: bool = False):
                para = doc.add_paragraph()
                try:
                    para.style = doc.styles[body_style]
                except KeyError:
                    pass
                if not is_code:
                    apply_body_paragraph_format(doc, para, style_name=body_style)
                run = para.add_run(text)
                apply_body_run_format(doc, run, style_name=body_style, is_code=is_code)
                if insert_first:
                    parent = after_element.getparent() if after_element is not None else para._element.getparent()
                    parent.insert(0, para._element)
                elif after_element is not None:
                    after_element.addnext(para._element)
                return para._element

            if position == "replace":
                anchor_para.clear()
                run = anchor_para.add_run(paragraphs_text[0])
                apply_body_run_format(doc, run, style_name=body_style, is_code=is_code)
                insert_after = anchor_para._element
                for extra_text in paragraphs_text[1:]:
                    insert_after = _styled_paragraph(extra_text, insert_after)
                result_msg = f"✅ 已替换段落 [{anchor_idx}] 的内容"
            else:
                insert_after_element = anchor_para._element
                remaining = paragraphs_text
                if position == "before":
                    prev = anchor_para._element.getprevious()
                    if prev is not None:
                        insert_after_element = prev
                    else:
                        insert_after_element = _styled_paragraph(
                            remaining[0], anchor_para._element, insert_first=True
                        )
                        remaining = remaining[1:]

                for text in remaining:
                    insert_after_element = _styled_paragraph(text, insert_after_element)

                pos_word = "后" if position == "after" else "前"
                result_msg = f"✅ 已在段落 [{anchor_idx}] \"{anchor_text[:30]}\" {pos_word}插入 {len(paragraphs_text)} 段内容"

            doc.save(str(file_path))
            self._notify_file_changed()
            size_kb = file_path.stat().st_size / 1024
            return f"{result_msg}\n文件大小: {size_kb:.1f} KB"

        except ImportError:
            return "Error: 缺少 python-docx 依赖，无法编辑模板"
        except Exception as e:
            logger.error(f"write_to_template 失败: {e}", exc_info=True)
            return f"Error: {e}"

    async def fill_template_table(
        self,
        table_index: int,
        content_json: str,
        start_row: int = 1,
        filename: str = "paper.docx",
        match_header: str = "",
    ) -> str:
        """
        按单元格填写模板里已有的表格，保留表格线和单元格原有字体。

        write_to_template 只改段落，改不了表格。课程报告里的测试用例表、
        评分表数据行必须走这个工具。

        Args:
            table_index: 表格序号，与 get_template_structure 中的「表格N」一致
            content_json: JSON 二维数组，每一行对应表格的一行单元格文本
            start_row: 从哪一行开始覆盖，默认 1（跳过表头）
            filename: 目标文件
            match_header: 可选，表头需包含的文字，用来核对没有填错表
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"

        try:
            rows = json.loads(content_json)
        except json.JSONDecodeError as exc:
            return f"Error: content_json 不是合法 JSON: {exc}"
        if not isinstance(rows, list) or not rows:
            return "Error: content_json 必须是非空二维数组，例如 [[\"1\",\"开始游戏\",...]]"
        if not all(isinstance(row, list) for row in rows):
            return "Error: content_json 每一行都必须是数组"

        try:
            from copy import deepcopy

            from docx import Document as PythonDocxDocument
            from docx.table import _Cell

            doc = PythonDocxDocument(str(file_path))
            if table_index < 0 or table_index >= len(doc.tables):
                return (
                    f"Error: 表格序号 {table_index} 不存在。"
                    f"当前文档共 {len(doc.tables)} 个表格。"
                )

            table = doc.tables[table_index]
            header = " | ".join(cell.text.strip() for cell in table.rows[0].cells)
            if match_header and match_header not in header:
                headers = [
                    f"表格{i}: {' | '.join(c.text.strip()[:20] for c in t.rows[0].cells)}"
                    for i, t in enumerate(doc.tables) if t.rows
                ]
                return (
                    f"Error: 表格{table_index} 表头不包含 \"{match_header}\"。\n"
                    f"当前表头: {header}\n" + "\n".join(headers)
                )

            col_count = len(table.columns) if table.rows else 0

            def _write_cell(cell: _Cell, text: str) -> None:
                paragraph = cell.paragraphs[0]
                r_pr = None
                for existing_run in paragraph.runs:
                    if existing_run._element.rPr is not None:
                        r_pr = deepcopy(existing_run._element.rPr)
                        break
                for extra in cell.paragraphs[1:]:
                    extra._element.getparent().remove(extra._element)
                paragraph.clear()
                run = paragraph.add_run(text)
                if r_pr is not None:
                    r_el = run._element
                    existing = r_el.rPr
                    if existing is not None:
                        r_el.remove(existing)
                    r_el.insert(0, r_pr)
                else:
                    apply_body_run_format(doc, run)

            def _clone_row():
                tbl = table._tbl
                last_tr = table.rows[-1]._tr
                new_tr = deepcopy(last_tr)
                tbl.append(new_tr)

            needed_rows = start_row + len(rows)
            while len(table.rows) < needed_rows:
                _clone_row()

            filled = 0
            for offset, row_values in enumerate(rows):
                row_idx = start_row + offset
                if row_idx >= len(table.rows):
                    break
                cells = table.rows[row_idx].cells
                for col_idx, value in enumerate(row_values[:col_count]):
                    _write_cell(cells[col_idx], "" if value is None else str(value))
                    filled += 1

            doc.save(str(file_path))
            self._notify_file_changed()
            return (
                f"✅ 已填写表格{table_index}（表头: {header[:40]}）"
                f"，从第 {start_row} 行起覆盖 {len(rows)} 行、共 {filled} 个单元格。"
            )
        except ImportError:
            return "Error: 缺少 python-docx 依赖，无法填写表格"
        except Exception as e:
            logger.error("fill_template_table 失败: %s", e, exc_info=True)
            return f"Error: {e}"

    async def repair_template_structure(self, filename: str = "paper.docx") -> str:
        """
        按上传模板恢复 Word 文档的标题骨架。

        这是一个受控修复工具，只修正与模板同序号的标题段落文本，
        不删除正文、不重排段落、不修改表格和图片。
        """
        file_path = self.workspace_dir / filename
        template_path = self.workspace_dir / ".system" / "_template_original.docx"
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"
        if not template_path.exists():
            return "Error: 当前工作空间没有模板文件，无法执行模板结构修复"

        try:
            from docx import Document as PythonDocxDocument

            doc = PythonDocxDocument(str(file_path))
            template = PythonDocxDocument(str(template_path))

            def heading_paragraphs(document):
                return [
                    para for para in document.paragraphs
                    if para.text.strip()
                    and para.style
                    and para.style.name.lower().startswith("heading")
                ]

            current_headings = heading_paragraphs(doc)
            template_headings = heading_paragraphs(template)
            if not template_headings:
                return "Error: 模板中未检测到标题段落"
            if len(current_headings) != len(template_headings):
                return (
                    "Error: 当前文档标题数量与模板不一致，无法安全自动修复。"
                    f"模板 {len(template_headings)} 个，当前 {len(current_headings)} 个。"
                )

            changed = []
            for index, (current, expected) in enumerate(zip(current_headings, template_headings), start=1):
                expected_text = expected.text.strip()
                current_text = current.text.strip()
                current_style = current.style.name if current.style else ""
                expected_style = expected.style.name if expected.style else ""
                if current_style != expected_style or current_text != expected_text:
                    current.clear()
                    current.style = expected.style
                    current.add_run(expected_text)
                    changed.append(f"{index}: {current_text} -> {expected_text}")

            if not changed:
                return "✅ 标题骨架已与模板一致，无需修复"

            doc.save(str(file_path))
            self._notify_file_changed()
            return "✅ 已按模板恢复标题骨架:\n" + "\n".join(changed[:20])

        except ImportError:
            return "Error: 缺少 python-docx 依赖，无法修复模板结构"
        except Exception as e:
            logger.error(f"repair_template_structure 失败: {e}", exc_info=True)
            return f"Error: {e}"

    async def get_template_structure(self, filename: str = "paper.docx") -> str:
        """
        获取模板文档的详细结构，包括段落索引、样式和内容预览。
        用于帮助定位 write_to_template 的 anchor_text。

        Args:
            filename: 目标文件名（默认 paper.docx）

        Returns:
            文档结构描述
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"

        try:
            from docx import Document as PythonDocxDocument
            doc = PythonDocxDocument(str(file_path))

            lines = [f"文档共 {len(doc.paragraphs)} 段, {len(doc.tables)} 个表格\n"]
            lines.append("--- 段落结构 ---")
            for i, para in enumerate(doc.paragraphs):
                text = para.text.strip()
                if not text:
                    continue
                style_name = para.style.name if para.style else "?"
                preview = text[:80] + ("..." if len(text) > 80 else "")
                lines.append(f"[{i}] [{style_name}] {preview}")

            if doc.tables:
                lines.append(f"\n--- 表格 ({len(doc.tables)}个) ---")
                for t_idx, table in enumerate(doc.tables):
                    rows = len(table.rows)
                    cols = len(table.columns) if table.rows else 0
                    first_cell = table.rows[0].cells[0].text[:40] if rows > 0 else ""
                    lines.append(f"表格{t_idx}: {rows}行×{cols}列, 首格=\"{first_cell}\"")
                    preview_rows = min(rows, 8)
                    for r_idx in range(preview_rows):
                        cells = [
                            table.rows[r_idx].cells[c].text.replace("\n", " ")[:24]
                            for c in range(cols)
                        ]
                        lines.append(f"  r{r_idx}: {cells}")
                    if rows > preview_rows:
                        lines.append(f"  ... 还有 {rows - preview_rows} 行")

            images = inventory_docx_images(file_path)
            if images:
                lines.append("")
                lines.append(format_image_inventory(images, title="图片").rstrip())

            lines.append("")
            lines.append(extract_style_fingerprint(file_path).format_report("当前文档样式").rstrip())

            return "\n".join(lines)

        except ImportError:
            return "Error: 缺少 python-docx 依赖"
        except Exception as e:
            return f"Error: {e}"

    async def inspect_document_styles(self, filename: str = "paper.docx") -> str:
        """
        查看指定 Word 的页面、样式定义、标题/正文样例。
        写模板前看模板，写完后看成品，用的是同一套档案。
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            template = self.workspace_dir / ".system" / "_template_original.docx"
            if filename != "paper.docx" or not template.exists():
                return f"Error: 文件不存在: {filename}"
            file_path = template
            filename = ".system/_template_original.docx"
        fingerprint = extract_style_fingerprint(file_path)
        return f"文档: {filename}\n" + fingerprint.format_report("样式档案")

    async def compare_document_styles(
        self,
        expected_filename: str = ".system/_template_original.docx",
        actual_filename: str = "paper.docx",
    ) -> str:
        """
        对照模板与成品的页面、页眉页脚和关键样式定义。
        写完后必须调用，确认成品没有把宋体/小四/边距写成另一套。
        """
        expected_path = self.workspace_dir / expected_filename
        actual_path = self.workspace_dir / actual_filename
        if not expected_path.exists():
            return f"Error: 对照基准不存在: {expected_filename}"
        if not actual_path.exists():
            return f"Error: 成品不存在: {actual_filename}"

        expected = extract_style_fingerprint(expected_path)
        actual = extract_style_fingerprint(actual_path)
        issues = compare_style_fingerprints(expected, actual)
        report = format_style_comparison(expected, actual, issues)
        if issues:
            return f"⚠️ 成品样式与模板不一致（{len(issues)} 项）\n{report}"
        return f"✅ 成品样式与模板一致\n{report}"

    async def extract_template_images(self, filename: str = "paper.docx") -> str:
        """
        将 Word 中的嵌入图片提取到 .system/docx_images/<文件名>/，
        并返回带附近文字的图片清单，方便后续识别和按位插图。
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"

        output_dir = self.workspace_dir / ".system" / "docx_images" / Path(filename).stem
        images = extract_docx_images(file_path, output_dir)
        if not images:
            return f"{filename} 中没有嵌入图片"

        lines = [
            f"✅ 已提取 {len(images)} 张图片到 {output_dir.relative_to(self.workspace_dir)}",
            format_image_inventory(images, title="提取结果").rstrip(),
        ]
        return "\n".join(lines)

    async def insert_image_to_template(
        self,
        image_path: str,
        anchor_text: str = "",
        position: str = "after",
        width_inches: float = 5.2,
        caption: str = "",
        filename: str = "paper.docx",
    ) -> str:
        """
        在模板指定段落附近插入图片，保留原有标题骨架和样式。

        Args:
            image_path: 工作区内图片路径，如 outputs/chart.png 或 .system/docx_images/paper/image1.png
            anchor_text: 定位段落文本；为空则追加到文档末尾
            position: after / before / replace
            width_inches: 图片显示宽度（英寸）
            caption: 可选图题，插入在图片下方并居中
            filename: 目标 Word 文件
        """
        file_path = self.workspace_dir / filename
        if not file_path.exists():
            return f"Error: 文件不存在: {filename}"

        source = Path(image_path)
        if not source.is_absolute():
            source = self.workspace_dir / image_path
        source = source.resolve()
        if not str(source).startswith(str(self.workspace_dir)):
            return f"Error: 图片路径超出工作空间: {image_path}"
        if not source.exists() or not source.is_file():
            return f"Error: 图片不存在: {image_path}"

        try:
            from docx import Document as PythonDocxDocument
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            from docx.shared import Inches

            doc = PythonDocxDocument(str(file_path))
            width = max(1.5, min(float(width_inches or 5.2), 6.3))

            def _style_name(para) -> str:
                return (para.style.name if para.style else "").strip().lower()

            def _is_toc_para(para) -> bool:
                style_name = _style_name(para)
                return style_name.startswith("toc") or "目录" in style_name

            anchor_para = None
            anchor_idx = None
            if anchor_text.strip():
                target = anchor_text.strip()
                for exact in (True, False):
                    for i, para in enumerate(doc.paragraphs):
                        text = para.text.strip()
                        if not text or _is_toc_para(para):
                            continue
                        if text == target if exact else target in text:
                            anchor_para = para
                            anchor_idx = i
                            break
                    if anchor_para is not None:
                        break
                if anchor_para is None:
                    available = [
                        f"[{i}] {p.text[:60]}"
                        for i, p in enumerate(doc.paragraphs)
                        if p.text.strip() and not _is_toc_para(p)
                    ][:30]
                    return (
                        f"Error: 未找到包含 \"{anchor_text}\" 的段落。\n"
                        f"可用段落（前30个非空）:\n" + "\n".join(available)
                    )
                if position == "replace" and _style_name(anchor_para).startswith("heading"):
                    return (
                        "Error: 禁止用图片替换模板标题段落。"
                        "请使用 position='after' 把图片插到标题或占位说明后面。"
                    )
            else:
                anchor_para = doc.paragraphs[-1] if doc.paragraphs else doc.add_paragraph()
                anchor_idx = len(doc.paragraphs) - 1
                position = "after"

            def _new_centered_paragraph():
                para = doc.add_paragraph()
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                return para

            def _place_after(after_element, para):
                after_element.addnext(para._element)
                return para._element

            if position == "replace":
                anchor_para.clear()
                anchor_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = anchor_para.add_run()
                run.add_picture(str(source), width=Inches(width))
                insert_after = anchor_para._element
            else:
                picture_para = _new_centered_paragraph()
                picture_para.add_run().add_picture(str(source), width=Inches(width))
                if position == "before":
                    prev = anchor_para._element.getprevious()
                    if prev is None:
                        parent = anchor_para._element.getparent()
                        parent.insert(0, picture_para._element)
                        insert_after = picture_para._element
                    else:
                        insert_after = _place_after(prev, picture_para)
                else:
                    insert_after = _place_after(anchor_para._element, picture_para)

            if caption.strip():
                cap_para = _new_centered_paragraph()
                run = cap_para.add_run(caption.strip())
                run.bold = True
                insert_after.addnext(cap_para._element)

            doc.save(str(file_path))
            self._notify_file_changed()
            size_kb = file_path.stat().st_size / 1024
            loc = f"段落 [{anchor_idx}]" if anchor_text.strip() else "文档末尾"
            return (
                f"✅ 已插入图片 {source.name} 到 {loc} "
                f"(宽 {width:.1f} 英寸)\n文件大小: {size_kb:.1f} KB"
            )
        except ImportError:
            return "Error: 缺少 python-docx 依赖，无法插入图片"
        except Exception as e:
            logger.error("insert_image_to_template 失败: %s", e, exc_info=True)
            return f"Error: {e}"
