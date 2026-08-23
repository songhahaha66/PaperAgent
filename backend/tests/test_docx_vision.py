import asyncio
import importlib.util
import os
import struct
import sys
import zlib
from pathlib import Path

import pytest
from docx import Document
from docx.shared import Inches

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx.oxml.ns import qn
from docx.shared import Pt

from ai_system.config.api_settings import load_env_api_settings, normalize_openai_base_url
from ai_system.core_tools.docx_images import extract_docx_images, inventory_docx_images
from ai_system.core_tools.docx_styles import compare_style_fingerprints, extract_style_fingerprint
from ai_system.core_tools.docx_tools import DocxTools
from ai_system.core_tools.vision_tools import VisionTools


def _load_template_contract():
    module_path = Path(__file__).resolve().parents[1] / "services/file_services/template_contract.py"
    spec = importlib.util.spec_from_file_location("template_contract_for_test", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module


_template_contract = _load_template_contract()
_build_contract = _template_contract._build_contract
_copy_template_to_workspace = _template_contract._copy_template_to_workspace


def _tiny_png(path: Path, color: bytes = b"\xff\x00\x00", width: int = 96, height: int = 64) -> Path:
    def chunk(tag: bytes, data: bytes) -> bytes:
        return (
            struct.pack(">I", len(data))
            + tag
            + data
            + struct.pack(">I", zlib.crc32(tag + data) & 0xFFFFFFFF)
        )

    raw = (b"\x00" + color * width) * height
    path.write_bytes(
        b"\x89PNG\r\n\x1a\n"
        + chunk(b"IHDR", struct.pack(">IIBBBBB", width, height, 8, 2, 0, 0, 0))
        + chunk(b"IDAT", zlib.compress(raw))
        + chunk(b"IEND", b"")
    )
    return path


def _docx_with_image(tmp_path: Path) -> Path:
    image = _tiny_png(tmp_path / "logo.png")
    doc_path = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("实验报告封面", style="Heading 1")
    doc.add_picture(str(image), width=Inches(1.2))
    doc.add_paragraph("1.2 实验结果", style="Heading 2")
    doc.add_paragraph("请在此插入收敛曲线图")
    doc.save(str(doc_path))
    return doc_path


def test_normalize_openai_base_url_appends_v1():
    assert normalize_openai_base_url("http://104.249.156.203:30080") == "http://104.249.156.203:30080/v1"
    assert normalize_openai_base_url("http://104.249.156.203:30080/v1") == "http://104.249.156.203:30080/v1"
    assert normalize_openai_base_url("https://api.openai.com/v1") == "https://api.openai.com/v1"


def test_inventory_and_extract_docx_images(tmp_path: Path):
    docx_path = _docx_with_image(tmp_path)
    images = inventory_docx_images(docx_path)

    assert len(images) == 1
    assert images[0].filename.lower().endswith((".png", ".jpeg", ".jpg"))
    assert "实验报告封面" in images[0].nearby_text or "实验结果" in images[0].nearby_text

    extracted = extract_docx_images(docx_path, tmp_path / "extracted")
    assert extracted[0].extracted_path
    assert Path(extracted[0].extracted_path).exists()


def test_word_contract_lists_embedded_images(tmp_path: Path):
    docx_path = _docx_with_image(tmp_path)
    contract = _build_contract(docx_path, "带图模板", "word")

    assert "图片骨架" in contract
    assert "模板样式档案" in contract
    assert "嵌入图片" in contract or "张嵌入图片" in contract
    assert "logo" in contract.lower() or "image" in contract.lower()


def test_copy_template_does_not_require_images(tmp_path: Path):
    template = tmp_path / "plain.docx"
    doc = Document()
    doc.add_paragraph("只有文字")
    doc.save(str(template))
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    _copy_template_to_workspace(template, workspace, "word")
    assert (workspace / "paper.docx").exists()


def test_insert_image_to_template_keeps_heading(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    chart = _tiny_png(workspace / "chart.png", color=b"\x00\x80\xff")

    doc = Document()
    doc.add_paragraph("1.2 实验结果", style="Heading 2")
    doc.add_paragraph("请在此插入收敛曲线图")
    doc.save(str(workspace / "paper.docx"))

    tools = DocxTools(str(workspace))
    result = asyncio.run(
        tools.insert_image_to_template(
            image_path="chart.png",
            anchor_text="请在此插入收敛曲线图",
            position="after",
            caption="图1 收敛曲线",
        )
    )

    assert "Error" not in result
    updated = Document(str(workspace / "paper.docx"))
    texts = [p.text for p in updated.paragraphs]
    assert texts[0] == "1.2 实验结果"
    assert "图1 收敛曲线" in texts
    assert inventory_docx_images(workspace / "paper.docx")


def test_insert_image_rejects_heading_replace(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    chart = _tiny_png(workspace / "chart.png")
    doc = Document()
    doc.add_paragraph("1.2 实验结果", style="Heading 2")
    doc.save(str(workspace / "paper.docx"))

    tools = DocxTools(str(workspace))
    result = asyncio.run(
        tools.insert_image_to_template(
            image_path=str(chart.name),
            anchor_text="1.2 实验结果",
            position="replace",
        )
    )

    assert "禁止用图片替换模板标题段落" in result
    assert Document(str(workspace / "paper.docx")).paragraphs[0].text == "1.2 实验结果"


def test_get_template_structure_lists_images(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    source = _docx_with_image(tmp_path)
    (workspace / "paper.docx").write_bytes(source.read_bytes())

    tools = DocxTools(str(workspace))
    structure = asyncio.run(tools.get_template_structure())
    extracted = asyncio.run(tools.extract_template_images())

    assert "图片" in structure
    assert "当前文档样式" in structure
    assert "已提取" in extracted
    assert (workspace / ".system" / "docx_images" / "paper").exists()


def test_analyze_image_without_api_returns_metadata(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    image = _tiny_png(workspace / "cover.png")
    monkeypatch.delenv("PAPERAGENT_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    tools = VisionTools(str(workspace))
    result = asyncio.run(tools.analyze_image("cover.png"))

    assert "cover.png" in result
    assert "未配置视觉模型" in result or "识别结果" in result


@pytest.mark.skipif(not os.getenv("PAPERAGENT_API_KEY"), reason="PAPERAGENT_API_KEY 未配置")
def test_live_vision_api_recognizes_color(tmp_path: Path):
    settings = load_env_api_settings()
    assert settings is not None
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    _tiny_png(workspace / "blue.png", color=b"\x1f\x4e\x79")

    tools = VisionTools(str(workspace))
    result = asyncio.run(
        tools.analyze_image("blue.png", question="这张图主要是什么颜色？只用一句话中文回答。")
    )

    assert "识别结果" in result
    assert "视觉识别失败" not in result
    assert any(token in result for token in ["蓝", "青", "深蓝"])


def _set_east_asia_style(doc: Document, style_name: str, font_name: str, size_pt: float) -> None:
    style = doc.styles[style_name]
    style.font.name = "Times New Roman"
    style.font.size = Pt(size_pt)
    r_pr = style.element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), font_name)
    r_fonts.set(qn("w:ascii"), "Times New Roman")


def test_style_fingerprint_reads_page_and_cjk_font(tmp_path: Path):
    path = tmp_path / "styled.docx"
    doc = Document()
    _set_east_asia_style(doc, "Normal", "宋体", 12)
    _set_east_asia_style(doc, "Heading 1", "黑体", 16)
    doc.add_paragraph("封面标题", style="Heading 1")
    doc.add_paragraph("这是一段正文")
    doc.save(str(path))

    fingerprint = extract_style_fingerprint(path)
    assert fingerprint.page.get("width_pt")
    assert fingerprint.styles["Normal"]["eastAsia"] == "宋体"
    assert fingerprint.styles["Heading 1"]["eastAsia"] == "黑体"
    assert fingerprint.styles["Normal"]["size_pt"] == 12


def test_compare_document_styles_detects_font_and_margin_drift(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = Document()
    _set_east_asia_style(template, "Normal", "宋体", 12)
    template.sections[0].left_margin = Pt(72)
    template.add_paragraph("第一章", style="Heading 1")
    template.save(str(system_dir / "_template_original.docx"))

    paper = Document()
    _set_east_asia_style(paper, "Normal", "微软雅黑", 10.5)
    paper.sections[0].left_margin = Pt(36)
    paper.add_paragraph("第一章", style="Heading 1")
    paper.save(str(workspace / "paper.docx"))

    expected = extract_style_fingerprint(system_dir / "_template_original.docx")
    actual = extract_style_fingerprint(workspace / "paper.docx")
    issues = compare_style_fingerprints(expected, actual)

    assert any("中文字体" in item for item in issues)
    assert any("左边距" in item for item in issues)

    tools = DocxTools(str(workspace))
    report = asyncio.run(tools.compare_document_styles())
    assert "样式与模板不一致" in report
    inspect = asyncio.run(tools.inspect_document_styles())
    assert "样式档案" in inspect


def test_write_to_template_inherits_normal_cjk_font(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    doc = Document()
    _set_east_asia_style(doc, "Normal", "宋体", 12)
    doc.add_paragraph("1.2 实验结果", style="Heading 2")
    doc.add_paragraph("请在此处填写")
    doc.save(str(workspace / "paper.docx"))

    tools = DocxTools(str(workspace))
    result = asyncio.run(
        tools.write_to_template(
            anchor_text="请在此处填写",
            content="这里是按模板样式写入的正文。",
            position="after",
        )
    )

    assert "Error" not in result
    updated = Document(str(workspace / "paper.docx"))
    texts = [p.text for p in updated.paragraphs]
    assert texts[:3] == ["1.2 实验结果", "请在此处填写", "这里是按模板样式写入的正文。"]
    inserted = next(p for p in updated.paragraphs if "按模板样式" in p.text)
    assert inserted.style.name == "Normal"
    east_asia = inserted.runs[0]._element.find(".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts")
    assert east_asia is not None
    assert east_asia.get(qn("w:eastAsia")) == "宋体"


def test_write_to_template_inherits_body_sample_when_normal_has_no_font(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    doc = Document()
    doc.styles["Normal"].paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    heading = doc.add_paragraph("1.1.1 内容简介", style="Heading 3")
    heading.runs[0].font.name = "黑体"
    r_pr = heading.runs[0]._element.get_or_add_rPr()
    r_fonts = r_pr.get_or_add_rFonts()
    r_fonts.set(qn("w:eastAsia"), "黑体")
    blank = doc.add_paragraph()
    blank.paragraph_format.line_spacing = 1.25
    sample = doc.add_paragraph("此部分为系统测试用例表，至少写5个以上测试用例。")
    sample.paragraph_format.line_spacing = 1.25
    run = sample.runs[0]
    run.font.size = Pt(12)
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run._element.rPr.rFonts.set(qn("w:ascii"), "宋体")
    doc.save(str(workspace / "paper.docx"))

    tools = DocxTools(str(workspace))
    result = asyncio.run(
        tools.write_to_template(
            anchor_text="1.1.1 内容简介",
            content="本实践围绕课程模板要求补充正文。",
            position="after",
        )
    )

    assert "Error" not in result
    updated = Document(str(workspace / "paper.docx"))
    inserted = next(p for p in updated.paragraphs if "课程模板要求" in p.text)
    assert inserted.paragraph_format.line_spacing == 1.25
    east_asia = inserted.runs[0]._element.find(
        ".//{http://schemas.openxmlformats.org/wordprocessingml/2006/main}rFonts"
    )
    assert east_asia is not None
    assert east_asia.get(qn("w:eastAsia")) == "宋体"
    assert inserted.runs[0].font.size is not None
    assert abs(inserted.runs[0].font.size.pt - 12) < 0.2


def test_review_agent_blocks_finished_style_drift(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = Document()
    _set_east_asia_style(template, "Heading 1", "黑体", 16)
    template.add_paragraph("第一章 DDL", style="Heading 1")
    template.save(str(system_dir / "_template_original.docx"))

    paper = Document()
    _set_east_asia_style(paper, "Heading 1", "楷体", 22)
    paper.add_paragraph("第一章 DDL", style="Heading 1")
    paper.add_paragraph("已经写了很多正文内容")
    paper.save(str(workspace / "paper.docx"))
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 DDL | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )

    module_path = Path(__file__).resolve().parents[1] / "ai_system/core_agents/review_agent.py"
    spec = importlib.util.spec_from_file_location("review_agent_style_test", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)

    reviewer = module.ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")
    result = asyncio.run(reviewer.review("完成这个实验报告"))

    assert result.complete is False
    assert "Word模板结构验收未通过" in result.reason


def test_analyze_docx_layout_includes_styles_without_images(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    doc = Document()
    _set_east_asia_style(doc, "Normal", "宋体", 12)
    doc.add_paragraph("只有文字的模板")
    doc.save(str(workspace / "paper.docx"))
    monkeypatch.delenv("PAPERAGENT_API_KEY", raising=False)
    monkeypatch.delenv("OPENAI_API_KEY", raising=False)

    result = asyncio.run(VisionTools(str(workspace)).analyze_docx_layout())
    assert "样式档案" in result
    assert "没有嵌入图片" in result
