from pathlib import Path
import sys
import asyncio
import importlib.util
import importlib
import subprocess
from unittest.mock import Mock
from types import SimpleNamespace

sys.path.insert(0, str(Path(__file__).resolve().parents[1]))

from docx import Document
from docx.enum.style import WD_STYLE_TYPE

from services.file_services.template_contract import (
    CONTRACT_PATH,
    CONTRACT_FILENAME,
    META_FILENAME,
    STYLE_PROFILE_FILENAME,
    _build_contract,
    _copy_template_to_workspace,
    _prepare_template_workspace,
    analyze_and_store_template,
    apply_stored_template_analysis,
    delete_template_analysis,
    ensure_template_analysis,
    read_template_analysis,
)
from services.file_services.plan_reconciler import PlanReconciler
from ai_system.core_tools.file_tools import FileTools
from ai_system.core_tools.docx_tools import DocxTools


def _load_review_agent_class():
    module_path = Path(__file__).resolve().parents[1] / "ai_system/core_agents/review_agent.py"
    spec = importlib.util.spec_from_file_location("review_agent_for_test", module_path)
    module = importlib.util.module_from_spec(spec)
    assert spec and spec.loader
    spec.loader.exec_module(module)
    return module.ReviewAgent


def _load_work_routes_module():
    return importlib.import_module("routers.work_routes.work")


def test_markdown_contract_extracts_headings_and_format_requirements(tmp_path: Path):
    template = tmp_path / "template.md"
    template.write_text(
        "# 论文题目\n\n"
        "要求：标题宋体小三居中。\n\n"
        "## 摘要\n\n"
        "## 1. 引言\n",
        encoding="utf-8",
    )

    contract = _build_contract(template, "测试模板", "markdown")

    assert "L1 第1行: 论文题目" in contract
    assert "L2 第5行: 摘要" in contract
    assert "宋体小三居中" in contract
    assert "必须完整保留模板骨架" in contract


def test_word_template_is_copied_as_initial_paper_docx(tmp_path: Path):
    template = tmp_path / "template.docx"
    doc = Document()
    doc.add_paragraph("标题要求：宋体，小三，居中")
    doc.add_paragraph("摘要")
    doc.save(str(template))

    workspace = tmp_path / "workspace"
    workspace.mkdir()

    _copy_template_to_workspace(template, workspace, "word")
    contract = _build_contract(template, "Word模板", "word")

    assert (workspace / "paper.docx").exists()
    assert (workspace / "paper.docx").read_bytes() == template.read_bytes()
    assert "标题要求" in contract
    assert "宋体" in contract
    assert "小三" in contract


def test_analyze_and_store_markdown_template_persists_contract(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(
        "services.file_services.template_contract.get_templates_path",
        lambda: tmp_path,
    )
    source = tmp_path / "course.md"
    source.write_text("# 题目\n\n要求：宋体小三居中。\n\n## 摘要\n", encoding="utf-8")

    contract = analyze_and_store_template(source, 7, "实验模板", "markdown")

    assert "宋体小三居中" in contract
    analysis = read_template_analysis(7)
    assert analysis["status"] == "ready"
    assert (tmp_path / ".analysis" / "7" / CONTRACT_FILENAME).exists()
    assert (tmp_path / ".analysis" / "7" / META_FILENAME).exists()


def test_prepare_workspace_reuses_upload_time_analysis(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(
        "services.file_services.template_contract.get_templates_path",
        lambda: tmp_path,
    )
    source = tmp_path / "course.md"
    source.write_text("# 题目\n\n要求：宋体小三居中。\n", encoding="utf-8")
    analyze_and_store_template(source, 11, "实验模板", "markdown")
    source.write_text("# 被改掉的模板\n", encoding="utf-8")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    contract = _prepare_template_workspace(source, workspace, 11, "实验模板", "markdown")

    assert "宋体小三居中" in contract
    assert "被改掉的模板" not in (workspace / CONTRACT_PATH).read_text(encoding="utf-8")
    assert (workspace / "paper.md").exists()


def test_prepare_workspace_reuses_word_style_profile(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(
        "services.file_services.template_contract.get_templates_path",
        lambda: tmp_path,
    )
    source = tmp_path / "course.docx"
    doc = Document()
    doc.add_paragraph("标题要求：宋体，小三，居中")
    doc.save(str(source))
    analyze_and_store_template(source, 12, "Word模板", "word")

    workspace = tmp_path / "workspace"
    workspace.mkdir()
    applied = apply_stored_template_analysis(12, workspace, "word")
    _copy_template_to_workspace(source, workspace, "word")

    assert applied is True
    assert (tmp_path / ".analysis" / "12" / STYLE_PROFILE_FILENAME).exists()
    assert (workspace / CONTRACT_PATH).exists()
    assert "宋体" in (workspace / CONTRACT_PATH).read_text(encoding="utf-8")
    assert (workspace / "paper.docx").exists()


def test_delete_template_analysis_removes_sidecar_dir(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(
        "services.file_services.template_contract.get_templates_path",
        lambda: tmp_path,
    )
    source = tmp_path / "course.md"
    source.write_text("# 题目\n", encoding="utf-8")
    analyze_and_store_template(source, 13, "实验模板", "markdown")
    assert (tmp_path / ".analysis" / "13").exists()

    delete_template_analysis(13)

    assert not (tmp_path / ".analysis" / "13").exists()


def test_ensure_template_analysis_backfills_old_template(tmp_path: Path, monkeypatch):
    monkeypatch.setattr(
        "services.file_services.template_contract.get_templates_path",
        lambda: tmp_path,
    )
    source = tmp_path / "old.md"
    source.write_text("# 题目\n\n要求：宋体小三居中。\n", encoding="utf-8")
    template = SimpleNamespace(id=21, file_path="old.md", name="旧模板", output_format="markdown")

    analysis = ensure_template_analysis(template)

    assert analysis["status"] == "ready"
    assert "宋体小三居中" in analysis["contract"]
    assert (tmp_path / ".analysis" / "21" / CONTRACT_FILENAME).exists()


def test_plan_system_loads_template_contract_into_plan_md_and_json(tmp_path: Path):
    workspace = tmp_path / "workspace"
    contract_path = workspace / CONTRACT_PATH
    contract_path.parent.mkdir(parents=True)
    contract_path.write_text("# 模板契约\n\n- 标题必须宋体小三居中\n", encoding="utf-8")

    reconciler = PlanReconciler(workspace)
    plan_md = reconciler.append_template_constraints("# 写作计划\n\n等待AI分析需求并制定写作计划...\n")
    structured = reconciler.build_from_markdown(plan_md)

    assert "## 模板强制约束" in plan_md
    assert "宋体小三居中" in plan_md
    assert structured["constraints"]["template_contract"].startswith("# 模板契约")
    assert structured["constraints"]["plan_markdown_synced"] is True


def test_write_to_template_skips_toc_and_preserves_heading(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    doc = Document()
    doc.styles.add_style("TOC 1", WD_STYLE_TYPE.PARAGRAPH)
    doc.add_paragraph("1.2.1 创建数据表", style="TOC 1")
    doc.add_paragraph("1.2.1 创建数据表", style="Heading 4")
    doc.add_paragraph("执行代码：")
    doc.save(str(workspace / "paper.docx"))

    tools = DocxTools(str(workspace))
    result = asyncio.run(
        tools.write_to_template(
            anchor_text="1.2.1 创建数据表",
            content="这里是实验说明。",
            position="after",
        )
    )

    assert "Error" not in result
    updated = Document(str(workspace / "paper.docx"))
    paragraphs = [p.text for p in updated.paragraphs]
    assert paragraphs[:4] == [
        "1.2.1 创建数据表",
        "1.2.1 创建数据表",
        "这里是实验说明。",
        "执行代码：",
    ]


def test_write_to_template_rejects_heading_replacement(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    doc = Document()
    doc.add_paragraph("1.2.1 创建数据表", style="Heading 4")
    doc.save(str(workspace / "paper.docx"))

    tools = DocxTools(str(workspace))
    result = asyncio.run(
        tools.write_to_template(
            anchor_text="1.2.1 创建数据表",
            content="错误替换标题",
            position="replace",
        )
    )

    assert "禁止替换模板标题段落" in result
    updated = Document(str(workspace / "paper.docx"))
    assert updated.paragraphs[0].text == "1.2.1 创建数据表"


def test_repair_template_structure_restores_heading_text(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = Document()
    template.add_paragraph("第一章 DDL", style="Heading 1")
    template.add_paragraph("1.2.1 创建数据表", style="Heading 4")
    template.add_paragraph("执行代码：")
    template.save(str(system_dir / "_template_original.docx"))

    paper = Document()
    paper.add_paragraph("第一章 DDL", style="Heading 1")
    paper.add_paragraph("**写在前面**：错误替换了标题", style="Heading 4")
    paper.add_paragraph("正文内容应保留")
    paper.save(str(workspace / "paper.docx"))

    tools = DocxTools(str(workspace))
    result = asyncio.run(tools.repair_template_structure())

    assert "已按模板恢复标题骨架" in result
    updated = Document(str(workspace / "paper.docx"))
    assert [p.text for p in updated.paragraphs] == [
        "第一章 DDL",
        "1.2.1 创建数据表",
        "正文内容应保留",
    ]
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 最终检查与完善 | ✅ 已完成 | 已验收 |\n",
        encoding="utf-8",
    )
    structured = PlanReconciler(workspace).build_from_markdown((workspace / "plan.md").read_text(encoding="utf-8"))
    assert structured["evidence"]["docx_template_issues"] == []


def test_review_agent_blocks_word_template_heading_drift(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = Document()
    template.add_paragraph("第一章 DDL", style="Heading 1")
    template.add_paragraph("1.2.1 创建数据表", style="Heading 4")
    template.add_paragraph("执行代码：")
    template.save(str(system_dir / "_template_original.docx"))

    paper = Document()
    paper.add_paragraph("第一章 DDL", style="Heading 1")
    paper.add_paragraph("**写在前面**：错误替换了标题", style="Heading 4")
    paper.add_paragraph("执行代码：")
    paper.save(str(workspace / "paper.docx"))

    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 DDL | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )

    ReviewAgent = _load_review_agent_class()
    reviewer = ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")
    result = asyncio.run(reviewer.review("完成这个实验报告"))

    assert result.complete is False
    assert "Word模板结构验收未通过" in result.reason


def test_review_agent_allows_dropping_template_instruction_parentheses(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = Document()
    template.add_paragraph(
        "2 项目开发报告（以下内容可根据自己的内容，进行替换更改结构，成稿后删除此括号，包含括号内容）",
        style="Heading 1",
    )
    template.add_paragraph(
        "2.4.2 核心代码（没有代码，可写制作流程，成稿后删除此括号，包含括号内容）",
        style="Heading 3",
    )
    template.save(str(system_dir / "_template_original.docx"))

    paper = Document()
    paper.add_paragraph("2 项目开发报告", style="Heading 1")
    paper.add_paragraph("2.4.2 核心代码", style="Heading 3")
    paper.save(str(workspace / "paper.docx"))

    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 2 项目开发报告 | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )

    tools = DocxTools(str(workspace))
    repair = asyncio.run(tools.repair_template_structure())
    assert "无需修复" in repair
    assert [p.text for p in Document(str(workspace / "paper.docx")).paragraphs] == [
        "2 项目开发报告",
        "2.4.2 核心代码",
    ]

    ReviewAgent = _load_review_agent_class()
    reviewer = ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")
    status = reviewer._read_word_status(workspace / "paper.docx", 1)
    assert "Word结构验收问题" not in status

    structured = PlanReconciler(workspace).build_from_markdown((workspace / "plan.md").read_text(encoding="utf-8"))
    assert structured["evidence"]["docx_template_issues"] == []


def test_review_agent_blocks_plan_blocked_status_without_llm(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    doc = Document()
    doc.add_paragraph("完整文档正文")
    doc.save(str(workspace / "paper.docx"))
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 最终检查与完善 | ❌ 阻塞 | 等待修复 |\n",
        encoding="utf-8",
    )

    ReviewAgent = _load_review_agent_class()
    reviewer = ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")
    result = asyncio.run(reviewer.review("完成这个实验报告"))

    assert result.complete is False
    assert "阻塞条目" in result.reason


def test_review_agent_ignores_status_words_in_template_constraints(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    doc = Document()
    doc.add_paragraph("完整文档正文")
    doc.save(str(workspace / "paper.docx"))
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 最终检查与完善 | ✅ 已完成 | 已验收 |\n\n"
        "<!-- template-constraints:start -->\n"
        "模板说明：这里可能出现 待写、进行中、阻塞 等普通文字。\n"
        "<!-- template-constraints:end -->\n",
        encoding="utf-8",
    )

    ReviewAgent = _load_review_agent_class()
    statuses = ReviewAgent._extract_plan_statuses((workspace / "plan.md").read_text(encoding="utf-8"))

    assert statuses == ["completed"]
    assert not ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")._deterministic_blockers(
        (workspace / "plan.md").read_text(encoding="utf-8"),
        "paper.docx: 1000 字节\nWord结构验收: 未发现模板结构问题",
    )


def test_review_agent_blocks_missing_plan_or_paper_without_llm(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()

    ReviewAgent = _load_review_agent_class()
    reviewer = ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")
    result = asyncio.run(reviewer.review("完成这个实验报告"))

    assert result.complete is False
    assert "plan.md 不存在" in result.reason
    assert "paper.docx 不存在" in result.reason


def test_review_agent_blocks_unparseable_plan_without_llm(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    doc = Document()
    doc.add_paragraph("完整文档正文")
    doc.save(str(workspace / "paper.docx"))
    (workspace / "plan.md").write_text("# 写作计划\n\n等待AI分析需求并制定写作计划...\n", encoding="utf-8")

    ReviewAgent = _load_review_agent_class()
    reviewer = ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")
    result = asyncio.run(reviewer.review("完成这个实验报告"))

    assert result.complete is False
    assert "未包含可解析的计划状态表" in result.reason


def test_plan_reconciler_blocks_completion_when_word_template_drifts(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = Document()
    template.add_paragraph("第一章 DDL", style="Heading 1")
    template.add_paragraph("1.2.1 创建数据表", style="Heading 4")
    template.save(str(system_dir / "_template_original.docx"))
    (system_dir / "template_contract.md").write_text("# 模板契约\n", encoding="utf-8")

    paper = Document()
    paper.add_paragraph("第一章 DDL", style="Heading 1")
    paper.add_paragraph("**写在前面**：错误替换了标题", style="Heading 4")
    paper.save(str(workspace / "paper.docx"))

    plan_md = (
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 DDL | ✅ 已完成 | 已写 |\n"
        "| 2 | 最终检查与完善 | ✅ 已完成 | 检查完成 |\n"
    )
    structured = PlanReconciler(workspace).build_from_markdown(plan_md)

    assert structured["stats"]["blocked"] == 1
    assert structured["stats"]["progress_percent"] < 100
    assert structured["evidence"]["docx_template_issues"]
    chapter = next(item for item in structured["items"] if item["title"] == "第一章 DDL")
    assert chapter["status"] == "pending"
    structure_item = next(item for item in structured["items"] if item["id"] == "task-template-structure")
    assert structure_item["status"] == "blocked"
    assert structure_item["title"] == "Word模板结构验收"
    assert structured["current_focus"]["title"] == "第一章 DDL"
    assert structure_item["raw_status"] == "template_validation_failed"
    assert structure_item["phase"] == "verify"


def test_plan_reconciler_syncs_metadata_progress_and_review_status(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "metadata.json").write_text(
        '{"work_id": "w1", "created_at": "2026-05-15", "status": "created", "progress": 0}',
        encoding="utf-8",
    )

    (workspace / "paper.md").write_text(
        "# 第一章\n\n" + "这一章已经写好了足够长的正文内容。" * 8 + "\n\n# 第二章\n\n",
        encoding="utf-8",
    )
    plan_md = (
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 | ✅ 已完成 | 已写 |\n"
        "| 2 | 第二章 | ⏳ 进行中 | 写作中 |\n"
    )
    reconciler = PlanReconciler(workspace)
    structured = reconciler.build_from_markdown(plan_md)
    reconciler.write_plan_json(structured)

    metadata = __import__("json").loads((workspace / "metadata.json").read_text(encoding="utf-8"))
    assert metadata["status"] == "running"
    assert metadata["progress"] == 50
    assert metadata["review_status"] == "in_progress"


def test_plan_reconciler_status_normalization_prioritizes_blocked(tmp_path: Path):
    reconciler = PlanReconciler(tmp_path)

    assert reconciler._normalize_plan_status("❌ 阻塞（原状态 ✅ 已完成）") == "blocked"
    assert reconciler._normalize_plan_status("⬜ 待写") == "pending"
    assert reconciler._normalize_plan_status("⏳ 进行中") == "in_progress"
    assert reconciler._normalize_plan_status("✅ 已完成") == "completed"


def test_plan_reconciler_syncs_metadata_even_when_plan_is_stable(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "metadata.json").write_text(
        '{"work_id": "w1", "created_at": "2026-05-15", "status": "created", "progress": 0}',
        encoding="utf-8",
    )
    plan_md = (
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 | ✅ 已完成 | 已写 |\n"
    )
    (workspace / "plan.md").write_text(plan_md, encoding="utf-8")
    (workspace / "paper.md").write_text("# 第一章\n\n" + "正文内容" * 80, encoding="utf-8")

    reconciler = PlanReconciler(workspace)
    first = reconciler.ensure_plan_json(sync_markdown=True)
    stale_metadata = __import__("json").loads((workspace / "metadata.json").read_text(encoding="utf-8"))
    stale_metadata["status"] = "created"
    stale_metadata["progress"] = 0
    stale_metadata.pop("review_status", None)
    (workspace / "metadata.json").write_text(
        __import__("json").dumps(stale_metadata, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )

    second = reconciler.ensure_plan_json(sync_markdown=True)
    metadata = __import__("json").loads((workspace / "metadata.json").read_text(encoding="utf-8"))

    assert first["stats"] == second["stats"]
    assert metadata["status"] == "completed"
    assert metadata["progress"] == 100
    assert metadata["review_status"] == "passed"


def test_plan_reconciler_does_not_complete_metadata_without_document(tmp_path: Path):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "metadata.json").write_text(
        '{"work_id": "w1", "created_at": "2026-05-15", "status": "created", "progress": 0}',
        encoding="utf-8",
    )
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )

    structured = PlanReconciler(workspace).ensure_plan_json(sync_markdown=True)
    metadata = __import__("json").loads((workspace / "metadata.json").read_text(encoding="utf-8"))

    assert structured["items"][0]["status"] == "pending"
    assert metadata["progress"] == 0
    assert metadata["status"] == "created"
    assert metadata["review_status"] == "pending"


def test_get_paper_status_reports_word_template_issues(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = Document()
    template.add_paragraph("第一章 DDL", style="Heading 1")
    template.add_paragraph("1.2.1 创建数据表", style="Heading 4")
    template.save(str(system_dir / "_template_original.docx"))
    (system_dir / "template_contract.md").write_text("# 模板契约\n", encoding="utf-8")

    paper = Document()
    paper.add_paragraph("第一章 DDL", style="Heading 1")
    paper.add_paragraph("**写在前面**：错误替换了标题", style="Heading 4")
    paper.save(str(workspace / "paper.docx"))
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 最终检查与完善 | ❌ 阻塞 | 等待修复 |\n",
        encoding="utf-8",
    )

    monkeypatch.setenv("WORKSPACE_DIR", str(workspace))
    status = FileTools().get_paper_status()

    assert "paper.docx 写作状态" in status
    assert "Word模板结构验收问题" in status
    assert "未完成/阻塞任务" in status


def test_ai_system_direct_import_uses_backend_config(tmp_path: Path):
    repo_backend = Path(__file__).resolve().parents[1]
    proc = subprocess.run(
        [
            sys.executable,
            "-c",
            (
                "import os; "
                f"os.environ['WORKSPACE_DIR'] = {str(tmp_path)!r}; "
                "from ai_system.core_tools.file_tools import FileTools; "
                "from config.paths import get_templates_path; "
                "print(FileTools.__name__, callable(get_templates_path))"
            ),
        ],
        cwd=str(repo_backend),
        text=True,
        capture_output=True,
        timeout=10,
    )

    assert proc.returncode == 0, proc.stderr
    assert "FileTools True" in proc.stdout


def test_get_work_metadata_reconciles_plan_before_reading(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "metadata.json").write_text(
        '{"work_id": "w1", "created_at": "2026-05-15", "status": "created", "progress": 0}',
        encoding="utf-8",
    )
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )
    (workspace / "paper.md").write_text("# 第一章\n\n" + "正文内容" * 80, encoding="utf-8")

    work_module = _load_work_routes_module()
    monkeypatch.setattr(work_module, "get_workspace_path", lambda work_id: workspace)
    monkeypatch.setattr(work_module.crud, "get_work", lambda db, work_id: Mock(created_by=7))

    metadata = __import__("asyncio").run(
        work_module.get_work_metadata("w1", db=object(), current_user=7)
    )

    assert metadata["status"] == "completed"
    assert metadata["progress"] == 100


def test_get_work_detail_syncs_status_from_workspace_metadata(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "metadata.json").write_text(
        '{"work_id": "w1", "created_at": "2026-05-15", "status": "created", "progress": 0}',
        encoding="utf-8",
    )
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )
    (workspace / "paper.md").write_text("# 第一章\n\n" + "正文内容" * 80, encoding="utf-8")

    work_module = _load_work_routes_module()
    work = Mock(work_id="w1", created_by=7, status="created", progress=0)
    monkeypatch.setattr(work_module, "get_workspace_path", lambda work_id: workspace)
    monkeypatch.setattr(work_module.crud, "get_work", lambda db, work_id: work)

    result = __import__("asyncio").run(work_module.get_work("w1", db=object(), current_user=7))

    assert result.status == "completed"
    assert result.progress == 100


def test_get_works_list_syncs_status_from_workspace_metadata(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    workspace.mkdir()
    (workspace / "metadata.json").write_text(
        '{"work_id": "w1", "created_at": "2026-05-15", "status": "created", "progress": 0}',
        encoding="utf-8",
    )
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 第一章 | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )
    (workspace / "paper.md").write_text("# 第一章\n\n" + "正文内容" * 80, encoding="utf-8")

    work_module = _load_work_routes_module()
    work = Mock(work_id="w1", created_by=7, status="created", progress=0)
    monkeypatch.setattr(work_module, "get_workspace_path", lambda work_id: workspace)
    monkeypatch.setattr(
        work_module.crud,
        "get_user_works",
        lambda db, current_user, skip, limit, status, search: {
            "works": [work],
            "total": 1,
            "page": 1,
            "size": 100,
        },
    )

    result = __import__("asyncio").run(
        work_module.get_works(db=object(), current_user=7)
    )

    assert result["works"][0].status == "completed"
    assert result["works"][0].progress == 100


def _homework_template_doc():
    doc = Document()
    doc.add_paragraph("信息技术实践与拓展实践报告")
    doc.add_paragraph("主    题")
    doc.add_paragraph("组    号            第     组")
    doc.add_paragraph("组长姓名")
    doc.add_paragraph("1 技术调研报告", style="Heading 1")
    doc.add_paragraph("1.1.1 内容简介", style="Heading 3")
    doc.add_paragraph("1.1.2 难点和解决办法", style="Heading 3")
    doc.add_paragraph("2.1 项目简介", style="Heading 2")
    doc.add_paragraph("图2.3 系统流程图")
    doc.add_paragraph("2.6 项目总结", style="Heading 2")
    doc.add_paragraph("总结项目总体完成情况、技术难点、解决方法以及不足等，不少于500字，填写后删除此项。")
    return doc


def test_plan_reconciler_does_not_complete_blank_word_template(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)

    template = _homework_template_doc()
    template.save(str(system_dir / "_template_original.docx"))
    _homework_template_doc().save(str(workspace / "paper.docx"))

    plan_md = (
        "| 序号 | 章节/任务 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 0 | 封面信息与提示清理 | ⏳ 进行中 | 填写主题、组号 |\n"
        "| 1 | 图表生成（CodeAgent） | ✅ 已完成 | 生成图2.1 |\n"
        "| 2 | 1 内容简介 | ✅ 已完成 | 学习总结 |\n"
        "| 3 | 2 难点和解决办法 | ✅ 已完成 | 旋转碰撞 |\n"
        "| 4 | 1 项目简介 | ⬜ 待写 | 项目背景 |\n"
    )
    structured = PlanReconciler(workspace).build_from_markdown(plan_md)
    by_title = {item["title"]: item for item in structured["items"]}

    def item_named(part: str):
        return next(item for item in structured["items"] if part in item["title"])

    assert by_title["封面信息与提示清理"]["status"] == "in_progress"
    assert by_title["图表生成（CodeAgent）"]["status"] == "pending"
    assert item_named("内容简介")["status"] == "pending"
    assert item_named("难点和解决办法")["status"] == "pending"
    assert item_named("项目简介")["status"] == "pending"
    assert structured["evidence"]["new_content_char_count"] == 0
    assert structured["evidence"]["new_image_count"] == 0
    assert structured["stats"]["completed"] == 0


def test_plan_reconciler_completes_section_after_new_body(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)
    _homework_template_doc().save(str(system_dir / "_template_original.docx"))

    paper = _homework_template_doc()
    paragraphs = list(paper.paragraphs)
    intro = next(p for p in paragraphs if "内容简介" in p.text)
    intro.add_run()
    intro._element.addnext(
        paper.add_paragraph("本节介绍 Java Swing 俄罗斯方块的游戏循环、方块旋转与消行规则，作为后续实现基础。" * 2)._element
    )
    paper.save(str(workspace / "paper.docx"))

    plan_md = (
        "| 序号 | 章节/任务 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 1 内容简介 | ⬜ 待写 | 学习总结 |\n"
        "| 2 | 2 难点和解决办法 | ✅ 已完成 | 尚未真正写 |\n"
    )
    structured = PlanReconciler(workspace).build_from_markdown(plan_md)
    def item_named(part: str):
        return next(item for item in structured["items"] if part in item["title"])

    assert item_named("内容简介")["status"] == "completed"
    assert item_named("难点和解决办法")["status"] == "pending"


def test_plan_reconciler_counts_only_new_output_images(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)
    _homework_template_doc().save(str(system_dir / "_template_original.docx"))
    _homework_template_doc().save(str(workspace / "paper.docx"))
    (workspace / "runs" / "run1").mkdir(parents=True)
    (workspace / "runs" / "run1" / "plot_1.png").write_bytes(b"not-a-real-png")

    plan_md = (
        "| 序号 | 章节/任务 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 图表生成（CodeAgent） | ✅ 已完成 | 生成图2.1 |\n"
    )
    structured = PlanReconciler(workspace).build_from_markdown(plan_md)
    assert structured["items"][0]["status"] == "pending"

    outputs = workspace / "outputs"
    outputs.mkdir()
    (outputs / "function_modules.png").write_bytes(b"not-a-real-png")
    structured = PlanReconciler(workspace).build_from_markdown(plan_md)
    assert structured["items"][0]["status"] == "completed"


def test_review_agent_blocks_unfilled_template_copy(tmp_path: Path):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)
    _homework_template_doc().save(str(system_dir / "_template_original.docx"))
    _homework_template_doc().save(str(workspace / "paper.docx"))
    (workspace / "plan.md").write_text(
        "| 序号 | 章节名 | 状态 | 说明 |\n"
        "|---|---|---|---|\n"
        "| 1 | 内容简介 | ✅ 已完成 | 已写 |\n",
        encoding="utf-8",
    )

    ReviewAgent = _load_review_agent_class()
    reviewer = ReviewAgent(llm=None, workspace_dir=str(workspace), output_mode="word")
    result = asyncio.run(reviewer.review("完成这个作业"))
    assert result.complete is False
    assert "未填写的模板骨架" in result.reason


def test_compare_paper_to_template_lists_empty_headings(tmp_path: Path, monkeypatch):
    workspace = tmp_path / "workspace"
    system_dir = workspace / ".system"
    system_dir.mkdir(parents=True)
    _homework_template_doc().save(str(system_dir / "_template_original.docx"))
    _homework_template_doc().save(str(workspace / "paper.docx"))
    monkeypatch.setenv("WORKSPACE_DIR", str(workspace))

    report = FileTools().compare_paper_to_template()
    assert "仍是未填写骨架" in report
    assert "内容简介" in report
    assert "空/仍是模板" in report
    assert "outputs 新图: 0" in report
