"""
Template contract extraction for paper generation.

The contract is a compact, model-readable summary of the uploaded template
bone structure and explicit formatting requirements. It is stored in the
workspace so both MainAgent and WriterAgent can enforce it.
"""

from __future__ import annotations

import json
import logging
import re
import shutil
from datetime import datetime, timezone
from pathlib import Path
from typing import Any, Iterable, Optional

from config.paths import get_templates_path
from ai_system.core_tools.docx_images import (
    extract_docx_images,
    format_image_inventory,
    inventory_docx_images,
)
from ai_system.core_tools.docx_styles import (
    PROFILE_PATH,
    extract_style_fingerprint,
    save_style_profile,
)

logger = logging.getLogger(__name__)

CONTRACT_PATH = ".system/template_contract.md"
ANALYSIS_DIRNAME = ".analysis"
CONTRACT_FILENAME = "contract.md"
STYLE_PROFILE_FILENAME = "style_profile.json"
IMAGES_DIRNAME = "images"
META_FILENAME = "meta.json"
WORKSPACE_IMAGES_DIR = Path(".system") / "docx_images" / "template"

FORMAT_KEYWORDS = (
    "宋体",
    "黑体",
    "仿宋",
    "楷体",
    "微软雅黑",
    "Times New Roman",
    "Arial",
    "小初",
    "初号",
    "小一",
    "一号",
    "小二",
    "二号",
    "小三",
    "三号",
    "小四",
    "四号",
    "小五",
    "五号",
    "字号",
    "字体",
    "加粗",
    "居中",
    "左对齐",
    "右对齐",
    "两端对齐",
    "行距",
    "页边距",
    "缩进",
    "首行",
)


def read_template_contract(workspace_dir: str | Path) -> str:
    path = Path(workspace_dir) / CONTRACT_PATH
    if not path.exists():
        return ""
    try:
        return path.read_text(encoding="utf-8")
    except Exception as exc:
        logger.warning("读取模板契约失败: %s", exc)
        return ""


def get_template_analysis_dir(template_id: int) -> Path:
    return get_templates_path() / ANALYSIS_DIRNAME / str(template_id)


def analyze_and_store_template(
    source_path: Path,
    template_id: int,
    template_name: str,
    output_mode: str,
) -> str:
    """Parse a template once at upload time and persist the analysis sidecars."""
    source_path = Path(source_path)
    analysis_dir = get_template_analysis_dir(template_id)
    if analysis_dir.exists():
        shutil.rmtree(analysis_dir)
    analysis_dir.mkdir(parents=True, exist_ok=True)

    error: Optional[str] = None
    contract = ""
    image_count = 0
    has_style_profile = False

    try:
        contract = _build_contract(source_path, template_name, output_mode)
        (analysis_dir / CONTRACT_FILENAME).write_text(contract, encoding="utf-8")
    except Exception as exc:
        error = str(exc)
        logger.warning("上传阶段抽取模板契约失败: %s", exc, exc_info=True)

    if source_path.suffix.lower() == ".docx":
        try:
            images = extract_docx_images(source_path, analysis_dir / IMAGES_DIRNAME)
            image_count = len(images or [])
        except Exception as exc:
            logger.warning("上传阶段提取模板图片失败: %s", exc)
            error = error or str(exc)
        try:
            fingerprint = extract_style_fingerprint(source_path)
            (analysis_dir / STYLE_PROFILE_FILENAME).write_text(
                json.dumps(fingerprint.to_dict(), ensure_ascii=False, indent=2),
                encoding="utf-8",
            )
            has_style_profile = True
        except Exception as exc:
            logger.warning("上传阶段保存模板样式档案失败: %s", exc)
            error = error or str(exc)

    meta = {
        "template_id": template_id,
        "template_name": template_name,
        "output_mode": output_mode,
        "status": "ready" if contract else "failed",
        "image_count": image_count,
        "has_style_profile": has_style_profile,
        "analyzed_at": datetime.now(timezone.utc).isoformat(),
        "error": error,
        "source_name": source_path.name,
    }
    (analysis_dir / META_FILENAME).write_text(
        json.dumps(meta, ensure_ascii=False, indent=2),
        encoding="utf-8",
    )
    return contract


def apply_stored_template_analysis(
    template_id: int,
    workspace_path: Path,
    output_mode: str = "",
) -> bool:
    """Copy upload-time analysis into a workspace. Returns True if a contract was applied."""
    del output_mode
    analysis_dir = get_template_analysis_dir(template_id)
    contract_file = analysis_dir / CONTRACT_FILENAME
    if not contract_file.exists():
        return False

    _write_contract(workspace_path, contract_file.read_text(encoding="utf-8"))

    style_src = analysis_dir / STYLE_PROFILE_FILENAME
    if style_src.exists():
        dest = Path(workspace_path) / PROFILE_PATH
        dest.parent.mkdir(parents=True, exist_ok=True)
        shutil.copyfile(style_src, dest)

    images_src = analysis_dir / IMAGES_DIRNAME
    if images_src.exists():
        dest = Path(workspace_path) / WORKSPACE_IMAGES_DIR
        if dest.exists():
            shutil.rmtree(dest)
        shutil.copytree(images_src, dest)
    return True


def read_template_analysis(template_id: int) -> dict[str, Any]:
    analysis_dir = get_template_analysis_dir(template_id)
    meta_file = analysis_dir / META_FILENAME
    contract_file = analysis_dir / CONTRACT_FILENAME
    meta: dict[str, Any] = {}
    if meta_file.exists():
        try:
            meta = json.loads(meta_file.read_text(encoding="utf-8"))
        except Exception:
            meta = {}
    contract = contract_file.read_text(encoding="utf-8") if contract_file.exists() else ""
    status = meta.get("status") or ("ready" if contract else "missing")
    return {
        "template_id": template_id,
        "status": status,
        "contract": contract,
        "image_count": int(meta.get("image_count") or 0),
        "has_style_profile": bool(
            meta.get("has_style_profile")
            or (analysis_dir / STYLE_PROFILE_FILENAME).exists()
        ),
        "analyzed_at": meta.get("analyzed_at"),
        "error": meta.get("error"),
    }


def ensure_template_analysis(template) -> dict[str, Any]:
    """Return stored analysis, backfilling old templates that were uploaded before this pipeline."""
    analysis_dir = get_template_analysis_dir(template.id)
    if not (analysis_dir / CONTRACT_FILENAME).exists() and getattr(template, "file_path", None):
        source_path = get_templates_path() / template.file_path
        if source_path.exists():
            analyze_and_store_template(
                source_path,
                template.id,
                template.name,
                getattr(template, "output_format", None) or "markdown",
            )
    return read_template_analysis(template.id)


def delete_template_analysis(template_id: int) -> None:
    analysis_dir = get_template_analysis_dir(template_id)
    if analysis_dir.exists():
        shutil.rmtree(analysis_dir)


def _prepare_template_workspace(
    source_path: Path,
    workspace_path: Path,
    template_id: int,
    template_name: str,
    output_mode: str,
) -> str:
    """Reuse upload-time analysis when present; otherwise parse now and persist for later workspaces."""
    analysis_dir = get_template_analysis_dir(template_id)
    if not (analysis_dir / CONTRACT_FILENAME).exists():
        analyze_and_store_template(source_path, template_id, template_name, output_mode)

    applied = apply_stored_template_analysis(template_id, workspace_path, output_mode)
    _copy_template_to_workspace(source_path, workspace_path, output_mode)
    if applied:
        return read_template_contract(workspace_path)

    contract = _build_contract(source_path, template_name, output_mode)
    _write_contract(workspace_path, contract)
    _capture_template_assets(source_path, workspace_path, output_mode)
    return contract


def initialize_template_contract(
    workspace_path: Path,
    template_id: Optional[int],
    output_mode: str,
) -> str:
    """
    Copy the uploaded template into the workspace and write a template contract.

    Prefers the analysis produced at upload time. Returns the contract text.
    Failures are logged and result in an empty contract so workspace creation
    can continue.
    """
    if template_id is None:
        return ""

    try:
        from database.database import SessionLocal
        from models.models import PaperTemplate

        db = SessionLocal()
        try:
            template = db.query(PaperTemplate).filter(PaperTemplate.id == template_id).first()
            if not template or not template.file_path:
                return ""
            source_path = get_templates_path() / template.file_path
            if not source_path.exists():
                logger.warning("模板文件不存在: %s", source_path)
                return ""
            return _prepare_template_workspace(
                source_path,
                workspace_path,
                template.id,
                template.name,
                output_mode,
            )
        finally:
            db.close()
    except Exception as exc:
        logger.warning("初始化模板契约失败: %s", exc, exc_info=True)
        return ""


def _copy_template_to_workspace(source_path: Path, workspace_path: Path, output_mode: str) -> None:
    if output_mode == "markdown":
        target = workspace_path / "paper.md"
        if not target.exists():
            shutil.copyfile(source_path, target)
    elif output_mode == "word":
        target = workspace_path / "paper.docx"
        if not target.exists():
            shutil.copyfile(source_path, target)
        backup = workspace_path / ".system" / "_template_original.docx"
        backup.parent.mkdir(parents=True, exist_ok=True)
        if not backup.exists():
            shutil.copyfile(source_path, backup)


def _capture_template_assets(source_path: Path, workspace_path: Path, output_mode: str) -> None:
    if output_mode != "word" or source_path.suffix.lower() != ".docx":
        return
    try:
        extract_docx_images(source_path, workspace_path / ".system" / "docx_images" / "template")
    except Exception as exc:
        logger.warning("提取模板图片失败: %s", exc)
    try:
        save_style_profile(workspace_path, extract_style_fingerprint(source_path))
    except Exception as exc:
        logger.warning("保存模板样式档案失败: %s", exc)


def _write_contract(workspace_path: Path, contract: str) -> None:
    path = workspace_path / CONTRACT_PATH
    path.parent.mkdir(parents=True, exist_ok=True)
    path.write_text(contract, encoding="utf-8")


def _build_contract(source_path: Path, template_name: str, output_mode: str) -> str:
    suffix = source_path.suffix.lower()
    if suffix == ".md":
        content = source_path.read_text(encoding="utf-8", errors="ignore")
        body = _markdown_contract(content)
    elif suffix == ".docx":
        body = _docx_contract(source_path)
    else:
        body = f"- 未支持的模板类型: {suffix}\n"

    return (
        "# 模板契约\n\n"
        f"- 模板名称: {template_name}\n"
        f"- 模板文件: {source_path.name}\n"
        f"- 输出模式: {output_mode}\n\n"
        "## 强制规则\n"
        "- 必须完整保留模板骨架中的标题层级、章节顺序、占位栏位、表格位置和显式说明。\n"
        "- 只能填充、扩展模板要求的内容；不要删除模板中的强制栏目，不要另起一套结构。\n"
        "- 模板中出现的字体、字号、对齐、行距、页边距等要求属于硬约束，例如“宋体”“小三”等必须遵循。\n"
        "- 如果模板要求与普通写作习惯冲突，优先执行模板要求。\n\n"
        f"{body}"
    )


def _markdown_contract(content: str) -> str:
    headings = []
    requirements = []
    for index, line in enumerate(content.splitlines(), start=1):
        stripped = line.strip()
        heading = re.match(r"^(#{1,6})\s+(.+)$", stripped)
        if heading:
            headings.append((len(heading.group(1)), heading.group(2), index))
        if _has_format_requirement(stripped):
            requirements.append((index, stripped))

    result = ["## Markdown 骨架"]
    if headings:
        for level, title, index in headings:
            result.append(f"- L{level} 第{index}行: {title}")
    else:
        result.append("- 未检测到 Markdown 标题；按原文顺序填充所有占位内容。")

    result.append("\n## 显式格式/写作要求")
    result.extend(_format_requirement_lines(requirements))
    return "\n".join(result) + "\n"


def _docx_contract(source_path: Path) -> str:
    try:
        from docx import Document
        from docx.shared import Pt

        doc = Document(str(source_path))
        paragraphs = []
        requirements = []
        style_lines = []

        for index, paragraph in enumerate(doc.paragraphs, start=1):
            text = paragraph.text.strip()
            if text:
                style_name = paragraph.style.name if paragraph.style else ""
                paragraphs.append((index, style_name, text))
                if _has_format_requirement(text):
                    requirements.append((index, text))

        for style in doc.styles:
            try:
                font = style.font
                attrs = []
                if font.name:
                    attrs.append(f"font={font.name}")
                if font.size:
                    attrs.append(f"size={_pt(font.size)}pt")
                if font.bold is not None:
                    attrs.append(f"bold={font.bold}")
                if attrs:
                    style_lines.append(f"- {style.name}: {', '.join(attrs)}")
            except Exception:
                continue

        result = ["## Word 骨架"]
        if paragraphs:
            for index, style_name, text in paragraphs[:80]:
                result.append(f"- 第{index}段 [{style_name}]: {_truncate(text)}")
            if len(paragraphs) > 80:
                result.append(f"- 其余 {len(paragraphs) - 80} 段按原模板顺序保留。")
        else:
            result.append("- 模板没有可抽取的正文段落；必须保留原 docx 样式与版式。")

        table_count = len(doc.tables)
        result.append(f"\n## 表格骨架\n- 模板包含 {table_count} 个表格，生成时必须保留相同表格位置和用途。")

        result.append("\n## 样式摘要")
        result.extend(style_lines[:60] if style_lines else ["- 未检测到显式样式；保留原 Word 模板样式。"])

        result.append("\n## 显式格式/写作要求")
        result.extend(_format_requirement_lines(requirements))
        result.append("")
        result.append(extract_style_fingerprint(source_path).format_report("模板样式档案").rstrip())
        result.append("")
        result.append(format_image_inventory(inventory_docx_images(source_path)).rstrip())
        return "\n".join(result) + "\n"
    except Exception as exc:
        logger.warning("抽取 Word 模板契约失败: %s", exc)
        return "## Word 骨架\n- 无法读取 docx 细节；必须以上传的 paper.docx 作为模板底稿继续写作。\n"


def _has_format_requirement(text: str) -> bool:
    return bool(text) and any(keyword in text for keyword in FORMAT_KEYWORDS)


def _format_requirement_lines(requirements: Iterable[tuple[int, str]]) -> list[str]:
    lines = [f"- 第{index}行/段: {_truncate(text, 160)}" for index, text in requirements]
    return lines or ["- 未检测到额外格式说明；仍需遵循模板已有样式和骨架。"]


def _truncate(text: str, limit: int = 120) -> str:
    text = re.sub(r"\s+", " ", text).strip()
    return text if len(text) <= limit else text[: limit - 1] + "…"


def _pt(value) -> str:
    try:
        return f"{value.pt:g}"
    except Exception:
        return str(value)
