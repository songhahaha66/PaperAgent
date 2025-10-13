import os
import json
import shutil
import zipfile
import tempfile
import logging
from pathlib import Path
from fastapi import HTTPException, status, UploadFile
from typing import List, Dict, Any, Optional
from .file_helper import FileHelper
from ..data_services.utils import handle_service_errors

logger = logging.getLogger(__name__)

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
    from docx.enum.style import WD_STYLE_TYPE
    DOCX_AVAILABLE = True
except ImportError:
    logger.warning("python-docx not available. DOCX export will be disabled.")
    DOCX_AVAILABLE = False

class WorkspaceFileService:
    def __init__(self):
        self.base_path = Path("../pa_data/workspaces")
        self.helper = FileHelper(self.base_path)

    def get_workspace_path(self, work_id: str) -> Path:
        """获取工作空间路径"""
        return self.base_path / work_id

    def ensure_workspace_exists(self, work_id: str) -> Path:
        """确保工作空间存在"""
        workspace_path = self.get_workspace_path(work_id)
        if not workspace_path.exists():
            # 只创建工作空间目录，不创建子目录
            workspace_path.mkdir(parents=True, exist_ok=True)
        return workspace_path

    def list_files(self, work_id: str, relative_path: str = "", recursive: bool = True) -> List[Dict[str, Any]]:
        """列出工作空间中的文件（新版本：按分类返回）

        Args:
            work_id: 工作ID
            relative_path: 相对路径，默认为根目录
            recursive: 是否递归遍历子目录，默认为True
        """
        # 直接调用新的分类方法，然后转换为平面列表格式
        categorized_files = self.list_files_by_category(work_id)
        result = []

        for category, files in categorized_files.items():
            for file_info in files:
                result.append(file_info)

        return result

    def list_files_by_category(self, work_id: str) -> Dict[str, List[Dict[str, Any]]]:
        """按分类列出工作空间中的文件

        Args:
            work_id: 工作ID

        Returns:
            包含五个分类的字典：{'code': [...], 'logs': [...], 'outputs': [...], 'papers': [...], 'attachments': [...]}
        """
        try:
            workspace_path = self.ensure_workspace_exists(work_id)

            # 确保三个分类文件夹存在（不包括papers）
            categories = {
                'code': workspace_path / 'code',
                'logs': workspace_path / 'logs',
                'outputs': workspace_path / 'outputs'
            }

            # 创建分类文件夹（如果不存在）
            for category_name, category_path in categories.items():
                if not category_path.exists():
                    category_path.mkdir(parents=True, exist_ok=True)

            result = {}

            # 处理三个分类文件夹
            for category_name, category_path in categories.items():
                files = []

                if category_path.exists():
                    # 扫描分类文件夹
                    for item in category_path.rglob('*'):
                        if item.is_file():
                            # 计算相对路径
                            rel_path = str(item.relative_to(category_path))
                            full_path = str(item.relative_to(workspace_path))

                            file_info = {
                                "name": item.name,
                                "type": "file",
                                "size": item.stat().st_size,
                                "modified": item.stat().st_mtime,
                                "path": full_path,
                                "category_path": rel_path,
                                "category": category_name
                            }
                            files.append(file_info)


                # 按名称排序
                files.sort(key=lambda x: x["name"].lower())
                result[category_name] = files

            # 特殊处理papers分类：扫描根目录的paper.md文件
            papers_files = []
            paper_md = workspace_path / 'paper.md'
            if paper_md.exists():
                file_info = {
                    "name": paper_md.name,
                    "type": "file",
                    "size": paper_md.stat().st_size,
                    "modified": paper_md.stat().st_mtime,
                    "path": "paper.md",
                    "category_path": "paper.md",
                    "category": "papers"
                }
                papers_files.append(file_info)

            result["papers"] = papers_files

            # 处理attachments分类：扫描attachment文件夹
            attachments_files = []
            attachment_dir = workspace_path / 'attachment'
            if attachment_dir.exists():
                for item in attachment_dir.rglob('*'):
                    if item.is_file():
                        # 计算相对路径
                        rel_path = str(item.relative_to(attachment_dir))
                        full_path = str(item.relative_to(workspace_path))

                        file_info = {
                            "name": item.name,
                            "type": "file",
                            "size": item.stat().st_size,
                            "modified": item.stat().st_mtime,
                            "path": full_path,
                            "category_path": rel_path,
                            "category": "attachments"
                        }
                        attachments_files.append(file_info)

            # 按名称排序
            attachments_files.sort(key=lambda x: x["name"].lower())
            result["attachments"] = attachments_files

            return result

        except HTTPException:
            raise
        except Exception as e:
            import logging
            logger = logging.getLogger(__name__)
            logger.error(f"按分类列出文件失败: {e}")
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to list files by category: {str(e)}"
            )

    
    @handle_service_errors()
    def read_file(self, work_id: str, file_path: str) -> Dict[str, Any]:
        """读取工作空间中的文件内容"""
        workspace_path = self.ensure_workspace_exists(work_id)
        target_file = workspace_path / file_path

        if not target_file.exists() or not target_file.is_file():
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")

        # 检测文件类型
        file_type = self.detect_file_type(file_path)

        if file_type == 'image':
            # 图片文件：返回base64编码的内容
            import base64
            if target_file.stat().st_size > 10 * 1024 * 1024:
                raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Image file too large to read")
            with open(target_file, 'rb') as f:
                content = base64.b64encode(f.read()).decode('utf-8')
            return {
                "type": "image",
                "content": content,
                "filename": target_file.name,
                "size": target_file.stat().st_size
            }

        elif file_type == 'text':
            # 文本文件：返回文件内容
            helper = FileHelper(workspace_path)
            content = helper.read_text(file_path)
            return {
                "type": "text",
                "content": content,
                "filename": target_file.name,
                "size": target_file.stat().st_size
            }

        else:  # binary
            # 二进制文件：返回元数据和下载信息
            import mimetypes
            mime_type, _ = mimetypes.guess_type(str(target_file))
            if not mime_type:
                mime_type = "application/octet-stream"

            return {
                "type": "binary",
                "filename": target_file.name,
                "size": target_file.stat().st_size,
                "mime_type": mime_type,
                "download_url": f"/api/workspace/{work_id}/download/{file_path}",
                "message": "Binary file - use download button to view"
            }

    def _is_image_file(self, file_path: str) -> bool:
        """判断文件是否为图片文件"""
        image_extensions = ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg', '.webp']
        return any(file_path.lower().endswith(ext) for ext in image_extensions)

    def detect_file_type(self, file_path: str) -> str:
        """检测文件类型：返回 'text', 'image', 'binary'"""

        # 图片文件扩展名
        image_extensions = {
            '.jpg', '.jpeg', '.png', '.gif', '.bmp', '.svg', '.webp',
            '.ico', '.tiff', '.tif'
        }

        # 文本文件扩展名
        text_extensions = {
            '.txt', '.md', '.py', '.js', '.ts', '.vue', '.html', '.css', '.scss', '.less',
            '.json', '.xml', '.yaml', '.yml', '.toml', '.ini', '.cfg', '.conf',
            '.c', '.cpp', '.cc', '.cxx', '.h', '.hpp', '.hxx',
            '.java', '.kt', '.scala', '.rs', '.go', '.php', '.rb', '.swift',
            '.sh', '.bash', '.zsh', '.fish', '.ps1', '.bat', '.cmd',
            '.sql', '.r', '.m', '.pl', '.lua', '.vim', '.dockerfile',
            '.gitignore', '.gitattributes', '.editorconfig', '.eslintrc', '.prettierrc',
            '.log', '.out', '.err', '.debug', '.trace'
        }

        # 二进制文件扩展名（常见文档、压缩包、可执行文件等）
        binary_extensions = {
            '.pdf', '.doc', '.docx', '.xls', '.xlsx', '.ppt', '.pptx',
            '.zip', '.rar', '.7z', '.tar', '.gz', '.bz2', '.xz',
            '.exe', '.msi', '.dmg', '.pkg', '.deb', '.rpm', '.apk',
            '.mp3', '.mp4', '.avi', '.mov', 'wmv', '.flv', '.mkv',
            '.ttf', '.otf', '.woff', '.woff2', '.eot',
            '.psd', '.ai', '.eps', '.sketch', '.fig'
        }

        # 获取文件扩展名
        ext = Path(file_path).suffix.lower()

        # 根据扩展名判断文件类型
        if ext in image_extensions:
            return 'image'
        elif ext in text_extensions:
            return 'text'
        elif ext in binary_extensions:
            return 'binary'
        else:
            # 未知扩展名，默认为二进制文件
            return 'binary'

    @handle_service_errors()
    def write_file(self, work_id: str, file_path: str, content: str) -> Dict[str, Any]:
        """写入文件到工作空间"""
        self.ensure_workspace_exists(work_id)
        helper = FileHelper(self.base_path / work_id)
        written_path = helper.write_text(Path(file_path).as_posix(), content)
        return {
            "message": "File written successfully",
            "path": Path(written_path).relative_to(self.base_path / work_id).as_posix(),
            "size": len(content)
        }

    @handle_service_errors()
    def upload_file(self, work_id: str, file_path: str, file: UploadFile) -> Dict[str, Any]:
        """上传文件到工作空间"""
        self.ensure_workspace_exists(work_id)
        helper = FileHelper(self.base_path / work_id)
        written_path = helper.write_bytes_stream(Path(file_path).as_posix(), file.file)
        size = Path(written_path).stat().st_size
        return {
            "message": "File uploaded successfully",
            "path": Path(written_path).relative_to(self.base_path / work_id).as_posix(),
            "size": size,
            "filename": file.filename
        }

    @handle_service_errors()
    def delete_file(self, work_id: str, file_path: str) -> Dict[str, str]:
        """删除工作空间中的文件或目录"""
        workspace_path = self.ensure_workspace_exists(work_id)
        # 防止删除整个工作空间
        if file_path.strip() in ("", ".", "/"):
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Cannot delete entire workspace")
        helper = FileHelper(workspace_path)
        helper.delete(Path(file_path).as_posix())
        return {"message": "File or directory deleted successfully"}

    @handle_service_errors()
    def create_directory(self, work_id: str, dir_path: str) -> Dict[str, str]:
        """在工作空间中创建目录"""
        workspace_path = self.ensure_workspace_exists(work_id)
        target_dir = workspace_path / dir_path
        if target_dir.exists():
            raise HTTPException(status_code=status.HTTP_400_BAD_REQUEST, detail="Directory already exists")
        target_dir.mkdir(parents=True, exist_ok=True)
        return {"message": "Directory created successfully"}

    @handle_service_errors()
    def get_file_info(self, work_id: str, file_path: str) -> Dict[str, Any]:
        """获取文件信息"""
        workspace_path = self.ensure_workspace_exists(work_id)
        target_file = workspace_path / file_path
        if not target_file.exists():
            raise HTTPException(status_code=status.HTTP_404_NOT_FOUND, detail="File not found")
        stat = target_file.stat()
        return {
            "name": target_file.name,
            "type": "directory" if target_file.is_dir() else "file",
            "size": stat.st_size if target_file.is_file() else None,
            "created": stat.st_ctime,
            "modified": stat.st_mtime,
            "path": str(target_file.relative_to(workspace_path))
        }

    def export_workspace(self, work_id: str) -> str:
        """导出工作空间为ZIP文件，返回临时文件路径"""
        try:
            workspace_path = self.ensure_workspace_exists(work_id)

            # 创建临时ZIP文件
            temp_dir = tempfile.mkdtemp()
            zip_filename = f"workspace_{work_id}.zip"
            zip_path = os.path.join(temp_dir, zip_filename)

            # 创建ZIP文件
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED) as zip_file:
                # 遍历工作空间目录
                for root, dirs, files in os.walk(workspace_path):
                    for file in files:
                        file_path = os.path.join(root, file)
                        # 计算相对路径
                        arcname = os.path.relpath(file_path, workspace_path)
                        zip_file.write(file_path, arcname)

                # 生成并添加docx文件
                try:
                    docx_content = self._generate_docx_from_paper(work_id)
                    if docx_content:
                        zip_file.writestr("paper.docx", docx_content)
                        logger.info(f"Successfully added paper.docx to workspace {work_id} export")
                    else:
                        logger.info(f"No paper.md found or docx generation failed for workspace {work_id}")
                except Exception as e:
                    logger.error(f"Failed to generate docx for workspace {work_id}: {str(e)}")
                    # 继续导出其他文件，不因docx生成失败而中断整个导出过程

                # 如果工作空间为空，添加一个空的README文件
                if not os.listdir(workspace_path):
                    zip_file.writestr("README.txt", "This workspace is empty.")

            return zip_path

        except HTTPException:
            raise
        except Exception as e:
            raise HTTPException(
                status_code=status.HTTP_500_INTERNAL_SERVER_ERROR,
                detail=f"Failed to export workspace: {str(e)}"
            )

    def _generate_docx_from_paper(self, work_id: str) -> Optional[bytes]:
        """生成paper.md对应的docx文件内容"""
        if not DOCX_AVAILABLE:
            logger.warning("python-docx not available, skipping docx generation")
            return None

        try:
            workspace_path = self.ensure_workspace_exists(work_id)
            paper_md_path = workspace_path / "paper.md"

            if not paper_md_path.exists():
                logger.info(f"paper.md not found in workspace {work_id}, skipping docx generation")
                return None

            # 读取paper.md内容
            with open(paper_md_path, 'r', encoding='utf-8') as f:
                markdown_content = f.read()

            # 转换为docx
            docx_content = self._convert_markdown_to_docx(markdown_content)
            return docx_content

        except Exception as e:
            logger.error(f"Failed to generate docx for workspace {work_id}: {str(e)}")
            return None

    def _convert_markdown_to_docx(self, markdown_content: str) -> bytes:
        """将Markdown内容转换为docx格式"""
        try:
            # 创建Word文档
            doc = Document()

            # 设置默认字体
            style = doc.styles['Normal']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(12)

            # 分割Markdown内容为行
            lines = markdown_content.split('\n')

            i = 0
            while i < len(lines):
                line = lines[i].strip()

                # 处理标题
                if line.startswith('# '):
                    heading = doc.add_heading(line[2:], level=1)
                    heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif line.startswith('## '):
                    heading = doc.add_heading(line[3:], level=2)
                elif line.startswith('### '):
                    heading = doc.add_heading(line[4:], level=3)
                elif line.startswith('#### '):
                    heading = doc.add_heading(line[5:], level=4)
                elif line.startswith('##### '):
                    heading = doc.add_heading(line[6:], level=5)
                # 处理空行
                elif line == '':
                    doc.add_paragraph()
                # 处理普通段落
                else:
                    # 简单的Markdown格式处理
                    processed_line = self._process_markdown_line(line)
                    paragraph = doc.add_paragraph(processed_line)

                i += 1

            # 保存到内存中的字节流
            from io import BytesIO
            doc_stream = BytesIO()
            doc.save(doc_stream)
            return doc_stream.getvalue()

        except Exception as e:
            logger.error(f"Failed to convert markdown to docx: {str(e)}")
            raise

    def _process_markdown_line(self, line: str) -> str:
        """处理单行Markdown格式，转换为纯文本"""
        # 简单处理，去除一些常见的Markdown标记
        line = line.replace('**', '').replace('*', '').replace('`', '')
        line = line.replace('[', '').replace(']', '').replace('(', '').replace(')', '')
        return line

# 创建全局实例
workspace_file_service = WorkspaceFileService()