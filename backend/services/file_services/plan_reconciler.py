"""
Workspace plan reconciliation.

The model can write a plan, but the workspace is the source of truth for
whether writing, charts, and final artifacts actually exist.
"""

from __future__ import annotations

import json
import re
import zipfile
from copy import deepcopy
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
from typing import Any, Dict, List, Optional

from .template_contract import CONTRACT_PATH


PLAN_PHASES = [
    {"id": "requirements", "title": "需求澄清", "description": "明确用户目标、输出格式和约束"},
    {"id": "design", "title": "方案设计", "description": "确定章节结构、数据/实验和产物形式"},
    {"id": "tasks", "title": "任务拆解", "description": "形成可执行、可追踪的章节和产物任务"},
    {"id": "implement", "title": "执行生成", "description": "逐项生成内容、代码、图表和最终文档"},
    {"id": "verify", "title": "验收检查", "description": "核对计划、正文、附件和最终产物"},
]

CONSTRAINTS_START = "<!-- template-constraints:start -->"
CONSTRAINTS_END = "<!-- template-constraints:end -->"


PLACEHOLDER_RE = re.compile(
    r"(成稿后删除[^。\n]*|填写后删除[^。\n]*|不少于\d+字[^。\n]*|"
    r"此部分为[^。\n]*|例如[:：][^。\n]*|待填写[^。\n]*|请填写[^。\n]*)"
)

MIN_NEW_SECTION_CHARS = 40


@dataclass
class PlanEvidence:
    document_text: str
    document_paths: List[str]
    document_char_count: int
    document_has_headings: bool
    docx_image_count: int
    docx_template_issues: List[str]
    output_images: List[str]
    generated_files: List[str]
    template_text: str = ""
    template_image_count: int = 0
    paper_blocks: Optional[List[tuple[str, str]]] = None
    template_blocks: Optional[List[tuple[str, str]]] = None
    new_content_char_count: int = 0
    new_image_count: int = 0

    @property
    def has_document(self) -> bool:
        return self.new_content_char_count >= MIN_NEW_SECTION_CHARS

    @property
    def has_output_image(self) -> bool:
        return self.new_image_count > 0

    def to_summary(self) -> Dict[str, Any]:
        return {
            "document_paths": self.document_paths,
            "document_char_count": self.document_char_count,
            "document_has_headings": self.document_has_headings,
            "docx_image_count": self.docx_image_count,
            "docx_template_issues": self.docx_template_issues,
            "output_images": self.output_images,
            "generated_files": self.generated_files[:30],
            "template_image_count": self.template_image_count,
            "new_content_char_count": self.new_content_char_count,
            "new_image_count": self.new_image_count,
        }


class PlanReconciler:
    """Build and refresh structured writing plans from local workspace state."""

    def __init__(self, workspace_path: Path | str):
        self.workspace_path = Path(workspace_path)

    def ensure_plan_json(self, sync_markdown: bool = True) -> Dict[str, Any]:
        plan_md_path = self.workspace_path / "plan.md"
        if plan_md_path.exists():
            plan_content = plan_md_path.read_text(encoding="utf-8")
        else:
            plan_content = "# 写作计划\n\n等待AI分析需求并制定写作计划...\n"

        previous_plan = self._read_existing_plan()
        previous_revision = int(previous_plan.get("revision", 0)) if previous_plan else 0
        structured_plan = self.build_from_markdown(
            plan_content,
            source="workspace_reconciled",
            revision=previous_revision if previous_plan else 1,
        )
        if previous_plan and self._stable_plan(previous_plan) == self._stable_plan(structured_plan):
            self._sync_metadata(previous_plan)
            return previous_plan

        structured_plan["revision"] = previous_revision + 1 if previous_plan else 1
        structured_plan["updated_at"] = datetime.now().isoformat()
        self.write_plan_json(structured_plan)
        if sync_markdown:
            self.write_plan_markdown(structured_plan)
        return structured_plan

    def build_from_markdown(
        self,
        plan_content: str,
        source: str = "update_plan_markdown",
        revision: Optional[int] = None,
    ) -> Dict[str, Any]:
        evidence = self._collect_evidence()
        items, title = self._parse_markdown_plan(plan_content)
        if self._is_placeholder_plan(items, plan_content) and evidence.has_document:
            items = self._synthesize_items_from_evidence(evidence)
        items = self._reconcile_items(items, evidence)
        structured_plan = self._assemble_plan(title, items, plan_content, source, evidence, revision=revision)
        return structured_plan

    def describe_progress_for_agent(self) -> str:
        """Facts for the model: template delta, empty headings, new images. No status guesses."""
        evidence = self._collect_evidence()
        lines = [
            "=== 相对模板的产物对照（只读，请据此判断哪些任务未完成） ===",
            f"文档字符数: {evidence.document_char_count}",
            f"相对模板新增正文(去掉占位句后): {evidence.new_content_char_count} 字",
            f"模板内嵌图: {evidence.template_image_count}，当前文档图: {evidence.docx_image_count}，outputs 新图: {len(evidence.output_images)}",
        ]
        if evidence.output_images:
            lines.append("outputs 文件: " + ", ".join(evidence.output_images[:20]))
        if evidence.template_text and evidence.new_content_char_count == 0:
            lines.append("注意: 当前正文相对模板没有新增内容，仍是未填写骨架。")
        headings = [text for kind, text in (evidence.paper_blocks or []) if kind == "heading"]
        if headings:
            lines.append("各标题下相对模板的正文:")
            for heading in headings:
                paper_body = self._section_body(evidence.paper_blocks or [], heading)
                template_body = self._section_body(evidence.template_blocks or [], heading)
                paper_n = len(self._meaningful_text(paper_body))
                template_n = len(self._meaningful_text(template_body))
                extra = paper_n - template_n
                mark = "有新增" if extra >= MIN_NEW_SECTION_CHARS else "空/仍是模板"
                lines.append(f"- {heading}: {mark}（当前{paper_n}字 / 模板{template_n}字）")
        issues = evidence.docx_template_issues or []
        if issues:
            lines.append("Word 结构差异:")
            lines.extend(f"- {issue}" for issue in issues[:5])
        return "\n".join(lines)

    def write_plan_json(self, structured_plan: Dict[str, Any]) -> None:
        plan_json_path = self.workspace_path / "plan.json"
        plan_json_path.write_text(json.dumps(structured_plan, ensure_ascii=False, indent=2), encoding="utf-8")
        self._sync_metadata(structured_plan)

    def write_plan_markdown(self, structured_plan: Dict[str, Any]) -> None:
        lines = [f"# {structured_plan.get('title') or '写作计划'}", ""]
        lines.extend([
            "| 序号 | 章节/任务 | 状态 | 说明 |",
            "|------|-----------|------|------|",
        ])
        for item in structured_plan.get("items", []):
            status_symbol = {
                "completed": "✅ 已完成",
                "in_progress": "⏳ 进行中",
                "blocked": "❌ 阻塞",
                "pending": "⬜ 待写",
            }.get(item.get("status"), "⬜ 待写")
            description = (item.get("description") or "").replace("\n", " ").strip()
            lines.append(f"| {item.get('order', '')} | {item.get('title', '')} | {status_symbol} | {description} |")

        current = structured_plan.get("current_focus")
        lines.append("")
        if current:
            lines.append(f"**当前阶段**: {current.get('title', '待执行')}")
        else:
            lines.append("**当前阶段**: 已完成")
        lines.append("")
        lines.append(f"**计划进度**: {structured_plan.get('stats', {}).get('progress_percent', 0)}%")
        template_contract = structured_plan.get("constraints", {}).get("template_contract") or self._read_template_contract()
        if template_contract:
            lines.extend(["", self._format_template_constraints(template_contract)])
        (self.workspace_path / "plan.md").write_text("\n".join(lines) + "\n", encoding="utf-8")

    def append_template_constraints(self, plan_content: str) -> str:
        """Attach template hard constraints to plan.md content without duplication."""
        clean_content = self._strip_template_constraints(plan_content).rstrip()
        template_contract = self._read_template_contract()
        if not template_contract:
            return clean_content + "\n"
        return clean_content + "\n\n" + self._format_template_constraints(template_contract) + "\n"

    def _parse_markdown_plan(self, plan_content: str) -> tuple[List[Dict[str, Any]], str]:
        plan_content = self._strip_template_constraints(plan_content)
        title = "写作计划"
        for line in plan_content.splitlines():
            stripped = line.strip()
            if stripped.startswith("#"):
                title = stripped.lstrip("#").strip() or title
                break

        rows: List[List[str]] = []
        for line in plan_content.splitlines():
            stripped = line.strip()
            if not stripped.startswith("|") or "---" in stripped:
                continue
            cells = self._split_markdown_table_row(stripped)
            if len(cells) >= 3:
                rows.append(cells)

        headers = rows[0] if rows else []
        data_rows = rows[1:] if rows else []

        def find_col(names: List[str], fallback: int) -> int:
            lowered = [h.lower() for h in headers]
            for idx, header in enumerate(lowered):
                if any(name in header for name in names):
                    return idx
            return fallback

        order_col = find_col(["序号", "order", "编号", "id"], 0)
        title_col = find_col(["章节", "任务", "名称", "title", "section", "step"], 1)
        status_col = find_col(["状态", "status"], 2)
        description_col = find_col(["说明", "描述", "description", "备注"], 3)

        items: List[Dict[str, Any]] = []
        for index, row in enumerate(data_rows, start=1):
            if not row or all(not cell for cell in row):
                continue
            raw_order = row[order_col] if order_col < len(row) else str(index)
            raw_title = row[title_col] if title_col < len(row) else f"任务 {index}"
            raw_status = row[status_col] if status_col < len(row) else ""
            raw_description = row[description_col] if description_col < len(row) else ""
            raw_status, raw_description = self._repair_status_description(raw_status, raw_description)
            order_match = re.search(r"\d+", raw_order)
            order = int(order_match.group()) if order_match else index
            status = self._normalize_plan_status(raw_status)
            items.append({
                "id": f"task-{order}",
                "order": order,
                "title": re.sub(r"^\s*\d+[\.\、]\s*", "", raw_title).strip() or f"任务 {order}",
                "status": status,
                "status_label": self._plan_status_label(status),
                "description": raw_description.strip(),
                "phase": "write",
                "depends_on": [f"task-{order - 1}"] if order > 1 else [],
                "raw_status": raw_status.strip(),
            })

        if not items:
            summary = plan_content.strip()
            if summary:
                items.append({
                    "id": "task-1",
                    "order": 1,
                    "title": summary.splitlines()[0].lstrip("#").strip()[:80] or "等待制定计划",
                    "status": "pending",
                    "status_label": self._plan_status_label("pending"),
                    "description": summary[:300],
                    "phase": "plan",
                    "depends_on": [],
                    "raw_status": "",
                })

        return sorted(items, key=lambda item: item["order"]), title

    def _is_placeholder_plan(self, items: List[Dict[str, Any]], plan_content: str) -> bool:
        if not items:
            return True
        content = plan_content.strip()
        if len(items) == 1 and any(token in content for token in ["等待AI", "等待 AI", "等待制定计划", "尚未制定"]):
            return True
        return False

    def _synthesize_items_from_evidence(self, evidence: PlanEvidence) -> List[Dict[str, Any]]:
        heading_titles = self._extract_document_headings(evidence.document_text)
        titles = ["论文结构"]
        titles.extend(heading_titles[:10])

        items: List[Dict[str, Any]] = []
        for order, title in enumerate(titles, start=1):
            items.append({
                "id": f"task-{order}",
                "order": order,
                "title": title,
                "status": "pending",
                "status_label": self._plan_status_label("pending"),
                "description": "由工作空间文档标题自动识别，需有相对模板的新增正文后才算完成",
                "phase": "write",
                "depends_on": [f"task-{order - 1}"] if order > 1 else [],
                "raw_status": "workspace_evidence",
                "status_source": "workspace_evidence",
            })
        return items

    def _extract_document_headings(self, document_text: str) -> List[str]:
        headings: List[str] = []
        patterns = [
            r"(?m)^#{1,6}\s+(.+)$",
            r"(?m)^\s*\d+(?:\.\d+)*[\.、]\s*(.{2,80})$",
        ]
        for pattern in patterns:
            for match in re.finditer(pattern, document_text):
                title = re.sub(r"\s+", " ", match.group(1)).strip()
                if title and title not in headings:
                    headings.append(title)
        return headings

    def _repair_status_description(self, raw_status: str, raw_description: str) -> tuple[str, str]:
        if raw_description.strip():
            return raw_status, raw_description
        match = re.match(r"^([✅⬜⏳❌]\s*(?:已完成|完成|待写|进行中|阻塞|失败)?)\s+(.+)$", raw_status.strip())
        if match:
            return match.group(1), match.group(2)
        return raw_status, raw_description

    def _reconcile_items(self, items: List[Dict[str, Any]], evidence: PlanEvidence) -> List[Dict[str, Any]]:
        if not items:
            return items

        reconciled = []
        for item in items:
            new_item = dict(item)
            evidence_status = self._infer_item_status(new_item, evidence)
            if evidence_status:
                new_item["status"] = evidence_status
                new_item["status_label"] = self._plan_status_label(evidence_status)
                new_item["status_source"] = "workspace_evidence"
            reconciled.append(new_item)

        if evidence.docx_template_issues:
            order = max((int(item.get("order", 0)) for item in reconciled), default=0) + 1
            existing = next((item for item in reconciled if item.get("id") == "task-template-structure"), None)
            payload = {
                "id": "task-template-structure",
                "order": existing.get("order", order) if existing else order,
                "title": "Word模板结构验收",
                "status": "blocked",
                "status_label": self._plan_status_label("blocked"),
                "description": "；".join(evidence.docx_template_issues[:3]),
                "phase": "verify",
                "depends_on": [f"task-{order - 1}"] if order > 1 else [],
                "raw_status": "template_validation_failed",
                "status_source": "template_validation",
            }
            if existing:
                existing.update(payload)
            else:
                reconciled.append(payload)

        return sorted(reconciled, key=lambda item: item["order"])

    def _infer_item_status(self, item: Dict[str, Any], evidence: PlanEvidence) -> Optional[str]:
        """Judge only by incremental products, never by keyword lists.

        If the task title maps to a document heading, require new body under
        that heading. If it does not map, only downgrade a completed mark when
        the whole paper still has no new text and no new images versus the
        template. The agent uses tools to inspect fields, figures, and covers.
        """
        if item.get("id") == "task-template-structure":
            return None
        current = item.get("status")
        if self._title_maps_to_heading(item, evidence):
            if self._section_has_new_body(item, evidence):
                return "completed"
            if current == "completed":
                return "pending"
            return None
        if current == "completed" and evidence.new_content_char_count == 0 and evidence.new_image_count == 0:
            return "pending"
        return None

    def _assemble_plan(
        self,
        title: str,
        items: List[Dict[str, Any]],
        source_markdown: str,
        source: str,
        evidence: PlanEvidence,
        revision: Optional[int] = None,
    ) -> Dict[str, Any]:
        stats = {
            "total": len(items),
            "completed": sum(1 for item in items if item["status"] == "completed"),
            "in_progress": sum(1 for item in items if item["status"] == "in_progress"),
            "blocked": sum(1 for item in items if item["status"] == "blocked"),
            "pending": sum(1 for item in items if item["status"] == "pending"),
        }
        stats["progress_percent"] = round((stats["completed"] / stats["total"]) * 100) if stats["total"] else 0

        current_focus = next((item for item in items if item["status"] == "in_progress"), None)
        if current_focus is None:
            current_focus = next((item for item in items if item["status"] == "pending"), None)
        if current_focus is None:
            current_focus = next((item for item in items if item["status"] == "blocked"), None)
        next_actions = [
            {
                "id": item["id"],
                "title": item["title"],
                "reason": "等待执行" if item["status"] == "pending" else "需要解除阻塞",
            }
            for item in items
            if item["status"] in {"pending", "blocked"}
        ][:3]

        previous_revision = 0
        plan_json_path = self.workspace_path / "plan.json"
        if plan_json_path.exists():
            try:
                previous_revision = int(json.loads(plan_json_path.read_text(encoding="utf-8")).get("revision", 0))
            except Exception:
                previous_revision = 0

        if revision is None:
            revision = previous_revision + 1

        return {
            "version": 1,
            "revision": revision,
            "title": title,
            "methodology": "spec-driven",
            "planning_mode": "dynamic",
            "phases": PLAN_PHASES,
            "items": items,
            "stats": stats,
            "current_focus": current_focus,
            "next_actions": next_actions,
            "source": source,
            "source_markdown": source_markdown,
            "constraints": {
                "template_contract": self._read_template_contract(),
                "plan_markdown_synced": bool(self._read_template_contract()),
            },
            "evidence": evidence.to_summary(),
            "updated_at": datetime.now().isoformat(),
        }

    def _read_template_contract(self) -> str:
        path = self.workspace_path / CONTRACT_PATH
        if not path.exists():
            return ""
        try:
            return path.read_text(encoding="utf-8").strip()
        except Exception:
            return ""

    def _strip_template_constraints(self, plan_content: str) -> str:
        pattern = re.compile(
            rf"\n?\s*{re.escape(CONSTRAINTS_START)}.*?{re.escape(CONSTRAINTS_END)}\s*",
            re.DOTALL,
        )
        return pattern.sub("\n", plan_content).strip() + "\n"

    def _format_template_constraints(self, template_contract: str) -> str:
        return (
            f"{CONSTRAINTS_START}\n"
            "## 模板强制约束\n\n"
            "以下内容由用户上传的模板骨架自动生成，模型写作、续写和验收时必须遵循。\n\n"
            f"{template_contract.strip()}\n"
            f"{CONSTRAINTS_END}"
        )

    def _read_existing_plan(self) -> Optional[Dict[str, Any]]:
        plan_json_path = self.workspace_path / "plan.json"
        if not plan_json_path.exists():
            return None
        try:
            return json.loads(plan_json_path.read_text(encoding="utf-8"))
        except Exception:
            return None

    def _stable_plan(self, plan: Dict[str, Any]) -> Dict[str, Any]:
        stable = deepcopy(plan)
        stable.pop("revision", None)
        stable.pop("updated_at", None)
        return stable

    def _sync_metadata(self, structured_plan: Dict[str, Any]) -> None:
        metadata_path = self.workspace_path / "metadata.json"
        if not metadata_path.exists():
            return

        try:
            metadata = json.loads(metadata_path.read_text(encoding="utf-8"))
        except Exception:
            metadata = {}

        stats = structured_plan.get("stats", {})
        progress = int(stats.get("progress_percent") or 0)
        has_started = any(
            int(stats.get(key) or 0) > 0
            for key in ("completed", "in_progress", "blocked")
        )
        is_complete = (
            progress == 100
            and int(stats.get("blocked") or 0) == 0
            and int(stats.get("pending") or 0) == 0
            and int(stats.get("in_progress") or 0) == 0
        )
        has_blocker = int(stats.get("blocked") or 0) > 0
        evidence = structured_plan.get("evidence", {})
        has_document = int(evidence.get("document_char_count") or 0) > 200

        metadata["progress"] = progress
        if is_complete and has_document:
            metadata["status"] = "completed"
            metadata["review_status"] = "passed"
            metadata.pop("review_reason", None)
        elif has_started:
            metadata["status"] = "running"
            metadata["review_status"] = "blocked" if has_blocker or is_complete else "in_progress"
            if has_blocker:
                blocker = next((
                    item for item in structured_plan.get("items", [])
                    if item.get("status") == "blocked"
                ), None)
                metadata["review_reason"] = (
                    blocker.get("description") if blocker else "计划存在阻塞项"
                )
            elif is_complete and not has_document:
                metadata["review_reason"] = "计划已完成但论文产物不存在或内容不足"
            else:
                metadata.pop("review_reason", None)
        else:
            metadata["status"] = metadata.get("status") or "created"
            metadata["review_status"] = "pending"
            metadata.pop("review_reason", None)

        metadata["updated_at"] = structured_plan.get("updated_at") or datetime.now().isoformat()
        metadata_path.write_text(json.dumps(metadata, ensure_ascii=False, indent=2), encoding="utf-8")

    def _collect_evidence(self) -> PlanEvidence:
        document_text_parts: List[str] = []
        document_paths: List[str] = []
        paper_blocks: List[tuple[str, str]] = []
        docx_image_count = 0
        docx_template_issues: List[str] = []

        paper_md = self.workspace_path / "paper.md"
        if paper_md.exists():
            text = paper_md.read_text(encoding="utf-8", errors="ignore")
            document_text_parts.append(text)
            document_paths.append("paper.md")
            paper_blocks.extend(self._parse_text_blocks(text))

        paper_docx = self.workspace_path / "paper.docx"
        if paper_docx.exists():
            text, image_count = self._read_docx(paper_docx)
            document_text_parts.append(text)
            document_paths.append("paper.docx")
            docx_image_count += image_count
            paper_blocks.extend(self._read_docx_blocks(paper_docx))
            docx_template_issues.extend(self._validate_docx_template(paper_docx))

        document_text = "\n".join(part for part in document_text_parts if part)

        template_text = ""
        template_blocks: List[tuple[str, str]] = []
        template_image_count = 0
        template_docx = self.workspace_path / ".system" / "_template_original.docx"
        if template_docx.exists():
            template_text, template_image_count = self._read_docx(template_docx)
            template_blocks = self._read_docx_blocks(template_docx)

        outputs_dir = self.workspace_path / "outputs"
        output_images = [
            str(path.relative_to(self.workspace_path))
            for path in outputs_dir.rglob("*")
            if path.is_file() and path.suffix.lower() in {".png", ".jpg", ".jpeg", ".gif", ".webp", ".svg"}
        ] if outputs_dir.exists() else []

        extra_docx_images = max(0, docx_image_count - template_image_count)
        new_image_count = len(output_images) + extra_docx_images

        generated_files = [
            str(path.relative_to(self.workspace_path))
            for path in self.workspace_path.rglob("*")
            if path.is_file()
            and not any(part in {".system", "__pycache__"} for part in path.relative_to(self.workspace_path).parts)
        ][:80] if self.workspace_path.exists() else []

        has_markdown_heading = bool(re.search(r"(?m)^#{1,6}\s+\S+", document_text))
        has_numbered_heading = bool(re.search(r"(?m)^\s*\d+(?:\.\d+)*[\.、]\s*\S+", document_text))
        new_content = self._meaningful_text(document_text)
        template_meaningful = self._meaningful_text(template_text)
        if template_meaningful and template_meaningful in new_content:
            incremental = new_content.replace(template_meaningful, "", 1)
        elif template_meaningful:
            incremental = new_content if new_content != template_meaningful else ""
        else:
            incremental = new_content

        return PlanEvidence(
            document_text=document_text,
            document_paths=document_paths,
            document_char_count=len(document_text.strip()),
            document_has_headings=has_markdown_heading or has_numbered_heading,
            docx_image_count=docx_image_count,
            docx_template_issues=docx_template_issues,
            output_images=sorted(output_images),
            generated_files=sorted(generated_files),
            template_text=template_text,
            template_image_count=template_image_count,
            paper_blocks=paper_blocks or None,
            template_blocks=template_blocks or None,
            new_content_char_count=len(incremental),
            new_image_count=new_image_count,
        )

    def _validate_docx_template(self, paper_docx: Path) -> List[str]:
        template_docx = self.workspace_path / ".system" / "_template_original.docx"
        if not template_docx.exists():
            return []
        try:
            paper_outline = self._docx_outline(paper_docx)
            template_outline = self._docx_outline(template_docx)
        except Exception as exc:
            return [f"无法解析Word结构: {exc}"]

        issues: List[str] = []
        if paper_outline["table_count"] != template_outline["table_count"]:
            issues.append(
                f"表格数量不一致（模板{template_outline['table_count']}，当前{paper_outline['table_count']}）"
            )
        from ai_system.core_tools.docx_styles import heading_outline_issues

        issues.extend(heading_outline_issues(template_outline["headings"], paper_outline["headings"]))
        markdown_heading_leaks = [
            text for _, text in paper_outline["headings"]
            if re.search(r"(^[*#`]+|[*#`]+$)", text)
        ]
        if markdown_heading_leaks:
            issues.append(f"标题中残留Markdown标记: {markdown_heading_leaks[:3]}")
        issues.extend(self._validate_docx_styles(template_docx, paper_docx))
        return issues

    @staticmethod
    def _validate_docx_styles(template_docx: Path, paper_docx: Path) -> List[str]:
        try:
            from ai_system.core_tools.docx_styles import compare_docx_styles

            return compare_docx_styles(template_docx, paper_docx)
        except Exception as exc:
            return [f"无法对照Word样式: {exc}"]

    def _docx_outline(self, docx_path: Path) -> Dict[str, Any]:
        import xml.etree.ElementTree as ET

        ns = {"w": "http://schemas.openxmlformats.org/wordprocessingml/2006/main"}
        with zipfile.ZipFile(docx_path) as archive:
            document_root = ET.fromstring(archive.read("word/document.xml"))
            style_names: Dict[str, str] = {}
            if "word/styles.xml" in archive.namelist():
                styles_root = ET.fromstring(archive.read("word/styles.xml"))
                for style in styles_root.findall(".//w:style", ns):
                    style_id = style.attrib.get(f"{{{ns['w']}}}styleId", "")
                    name_el = style.find("w:name", ns)
                    style_names[style_id] = (
                        name_el.attrib.get(f"{{{ns['w']}}}val", style_id)
                        if name_el is not None else style_id
                    )

        headings = []
        for para in document_root.findall(".//w:p", ns):
            text = "".join(t.text or "" for t in para.findall(".//w:t", ns)).strip()
            style_el = para.find("./w:pPr/w:pStyle", ns)
            style_id = style_el.attrib.get(f"{{{ns['w']}}}val", "") if style_el is not None else ""
            style_name = style_names.get(style_id, style_id)
            if style_name.lower().startswith("heading") and text:
                headings.append((style_name.lower(), text))

        return {
            "headings": headings,
            "table_count": len(document_root.findall(".//w:tbl", ns)),
        }

    def _read_docx(self, docx_path: Path) -> tuple[str, int]:
        try:
            from docx import Document

            document = Document(str(docx_path))
            paragraphs = [p.text for p in document.paragraphs if p.text.strip()]
            image_count = len(document.part._rels)
            image_count = sum(1 for rel in document.part._rels.values() if "image" in rel.reltype)
            return "\n".join(paragraphs), image_count
        except Exception:
            return self._read_docx_zip_fallback(docx_path)

    def _read_docx_zip_fallback(self, docx_path: Path) -> tuple[str, int]:
        try:
            with zipfile.ZipFile(docx_path) as archive:
                xml = archive.read("word/document.xml").decode("utf-8", errors="ignore")
                text = re.sub(r"<[^>]+>", "\n", xml)
                text = re.sub(r"\n+", "\n", text)
                media = [name for name in archive.namelist() if name.startswith("word/media/")]
                return text, len(media)
        except Exception:
            return "", 0

    def _normalize_plan_status(self, raw_status: str) -> str:
        text = (raw_status or "").strip().lower()
        if any(token in text for token in ["❌", "失败", "blocked", "阻塞", "error", "failed"]):
            return "blocked"
        if any(token in text for token in ["⬜", "待写", "pending", "todo"]):
            return "pending"
        if any(token in text for token in ["⏳", "进行", "progress", "doing", "current"]):
            return "in_progress"
        if any(token in text for token in ["✅", "完成", "done", "completed", "complete"]):
            return "completed"
        return "pending"

    def _plan_status_label(self, status: str) -> str:
        return {
            "completed": "已完成",
            "in_progress": "进行中",
            "blocked": "阻塞",
            "pending": "待写",
        }.get(status, "待写")

    def _split_markdown_table_row(self, line: str) -> List[str]:
        line = line.strip()
        if line.startswith("|"):
            line = line[1:]
        if line.endswith("|"):
            line = line[:-1]
        return [cell.strip() for cell in line.split("|")]

    def _has_any(self, text: str, keywords: List[str]) -> bool:
        return any(keyword.lower() in text for keyword in keywords)

    def _normalize_text(self, text: str) -> str:
        return re.sub(r"\s+", "", text).lower()

    def _meaningful_text(self, text: str) -> str:
        cleaned = PLACEHOLDER_RE.sub("", text or "")
        return re.sub(r"\s+", "", cleaned)

    def _heading_key(self, title: str) -> str:
        text = title or ""
        text = re.sub(r"（[^）]*）", "", text)
        text = re.sub(r"\([^)]*\)", "", text)
        text = re.sub(r"^[\d\.、\s]+", "", text)
        return self._normalize_text(text)

    def _title_maps_to_heading(self, item: Dict[str, Any], evidence: PlanEvidence) -> bool:
        key = self._heading_key(item.get("title") or "")
        if len(key) < 2:
            return False
        blocks = evidence.paper_blocks or self._parse_text_blocks(evidence.document_text)
        for kind, text in blocks:
            if kind != "heading":
                continue
            heading_key = self._heading_key(text)
            if heading_key == key or (len(key) >= 4 and (key in heading_key or heading_key in key)):
                return True
        return False

    def _section_has_new_body(self, item: Dict[str, Any], evidence: PlanEvidence) -> bool:
        paper_body = self._section_body(
            evidence.paper_blocks or self._parse_text_blocks(evidence.document_text),
            item.get("title") or "",
        )
        template_body = self._section_body(
            evidence.template_blocks or self._parse_text_blocks(evidence.template_text),
            item.get("title") or "",
        )
        paper_m = self._meaningful_text(paper_body)
        template_m = self._meaningful_text(template_body)
        if len(paper_m) < MIN_NEW_SECTION_CHARS:
            return False
        if not template_m:
            return True
        if template_m in paper_m:
            extra = paper_m.replace(template_m, "", 1)
            return len(extra) >= MIN_NEW_SECTION_CHARS
        return paper_m != template_m and len(paper_m) >= MIN_NEW_SECTION_CHARS

    def _section_body(self, blocks: List[tuple[str, str]], title: str) -> str:
        key = self._heading_key(title)
        if len(key) < 2:
            return ""
        capturing = False
        body: List[str] = []
        for kind, text in blocks:
            if kind == "heading":
                heading_key = self._heading_key(text)
                matched = heading_key == key or (
                    len(key) >= 4 and (key in heading_key or heading_key in key)
                )
                if capturing:
                    break
                if matched:
                    capturing = True
                continue
            if capturing:
                body.append(text)
        return "\n".join(body)

    def _parse_text_blocks(self, text: str) -> List[tuple[str, str]]:
        blocks: List[tuple[str, str]] = []
        for line in (text or "").splitlines():
            stripped = line.strip()
            if not stripped:
                continue
            if stripped.startswith("#"):
                blocks.append(("heading", stripped.lstrip("#").strip()))
            elif re.match(r"^\d+(?:\.\d+)*[\.、\s]+\S+", stripped):
                blocks.append(("heading", stripped))
            else:
                blocks.append(("body", stripped))
        return blocks

    def _read_docx_blocks(self, docx_path: Path) -> List[tuple[str, str]]:
        try:
            from docx import Document

            document = Document(str(docx_path))
            blocks: List[tuple[str, str]] = []
            for paragraph in document.paragraphs:
                text = (paragraph.text or "").strip()
                if not text:
                    continue
                style = (paragraph.style.name or "").lower() if paragraph.style else ""
                kind = "heading" if style.startswith("heading") or "标题" in style else "body"
                blocks.append((kind, text))
            return blocks
        except Exception:
            text, _ = self._read_docx_zip_fallback(docx_path)
            return self._parse_text_blocks(text)
